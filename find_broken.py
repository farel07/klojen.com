from docx import Document
doc = Document('laporan_pkl.docx')
NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# Find which paragraph uses which image relationship
for rel_id in ['rId29', 'rId30', 'rId31', 'rId32', 'rId33', 'rId35', 'rId36', 'rId43']:
    for i, p in enumerate(doc.paragraphs):
        for drawing in p._element.findall(f'.//{NS}drawing'):
            # Find blip element which has embed attribute
            blips = drawing.findall(f'.//{NS}blip')
            for blip in blips:
                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed == rel_id:
                    # Find nearest title paragraph (look backward)
                    title = ''
                    for j in range(i-1, max(0, i-5), -1):
                        t = doc.paragraphs[j].text.strip()
                        if t:
                            title = t
                            break
                    print(f'{rel_id} -> para [{i}], title: {title[:80]}')
