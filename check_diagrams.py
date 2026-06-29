from docx import Document
doc = Document('laporan_pkl.docx')
NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# Check all paragraphs for images around diagrams
for i in range(449, 555):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        drawings = p._element.findall(f'.//{NS}drawing')
        has_img = len(drawings) > 0
        if t or has_img:
            label = 'IMG' if has_img else 'TXT'
            text = t[:90] if t else '(empty)'
            print(f'[{i}] [{label}] {text}')
