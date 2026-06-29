from docx import Document
doc = Document('laporan_pkl.docx')

# Find paragraphs with [Gambar placeholder in content area (after para 400)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if i > 400 and ('[Gambar' in t or t.startswith('Alur:') or t.startswith('Alur ')):
        print(f'[{i}] {p.style.name!r}: {t[:150]}')

print('\n--- Also checking for text-based diagram descriptions ---')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if i > 400 and ('→' in t or '->>' in t or 'participant' in t.lower()):
        if len(t) > 50:
            print(f'[{i}] {t[:150]}')
