from docx import Document
from PIL import Image
import io, re
from lxml import etree

doc = Document('laporan_pkl.docx')

# Map each relationship ID to its image size
rel_sizes = {}
for rel_id, rel in doc.part.rels.items():
    if 'image' in rel.reltype:
        blob = rel.target_part.blob
        try:
            img = Image.open(io.BytesIO(blob))
            rel_sizes[rel_id] = img.size
        except:
            rel_sizes[rel_id] = (0, 0)

# Map each paragraph to its embedded images
para_images = {}
for i, p in enumerate(doc.paragraphs):
    xml_str = etree.tostring(p._element, encoding='unicode')
    embeds = re.findall(r'r:embed="([^"]+)"', xml_str)
    if embeds:
        para_images[i] = embeds

# List all diagram titles with their image status
diagram_titles = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    # Match specific diagram titles (not section headers)
    if any(t.startswith(prefix) for prefix in [
        'Activity Diagram —', 'Sequence Diagram —', 'Class Diagram —',
        'Use Case Diagram', 'State Machine'
    ]):
        diagram_titles.append((i, t))

print(f'Found {len(diagram_titles)} diagram titles\n')
for title_idx, title in diagram_titles:
    # Search forward for the nearest image
    found = False
    for j in range(title_idx + 1, min(title_idx + 10, len(doc.paragraphs))):
        if j in para_images:
            embeds = para_images[j]
            sizes = [rel_sizes.get(e, (0,0)) for e in embeds]
            min_pixels = min(s[0]*s[1] for s in sizes)
            if min_pixels > 10000:
                status = 'OK'
            elif min_pixels > 0:
                status = f'BROKEN {sizes}'
            else:
                status = f'UNKNOWN {sizes}'
            print(f'[{title_idx}] {title[:65]}')
            print(f'  -> para [{j}] {status}')
            found = True
            break
    if not found:
        print(f'[{title_idx}] {title[:65]}')
        print(f'  -> NO IMAGE (searched {title_idx+1}-{min(title_idx+10, len(doc.paragraphs))})')
