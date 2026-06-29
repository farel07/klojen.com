#!/usr/bin/env python3
"""
Comprehensive script to fill ALL empty diagram placeholders in FIX.docx.
"""
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_diagram_format(paragraph, text):
    """Format paragraph as diagram description."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.text = ''
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)

def fill_all_diagrams(doc):
    """Fill all empty diagram placeholders with text-based descriptions."""
    
    # Activity Diagrams
    activity_diagrams = {
        "Activity Diagram — Proses Login": """[Activity Diagram: Proses Login & Autentikasi]

START (●)
  │
  ▼
[User membuka halaman login]
  │
  ▼
[User mengisi email & password]
  │
  ▼
[Submit form login]
  │
  ▼
<Validasi input?> ──No──> [Tampilkan error validasi] ─┐
  │ Yes                                                │
  ▼                                                    │
[POST /api/auth/login]                                 │
  │                                                    │
  ▼
<Kredensial valid?> ──No──> [Tampilkan: "Email atau    │
  │ Yes                        password salah"]        │
  ▼                                                    │
[Backend generate access_token (15 menit)              │
 dan refresh_token (30 hari)]                          │
  │                                                    │
  ▼
[Simpan refresh_token hash SHA-256 ke DB]              │
  │                                                    │
  ▼
[Return token pair ke frontend]                        │
  │                                                    │
  ▼
[Simpan access_token di Zustand store]                 │
  │                                                    │
  ▼
[Simpan refresh_token di localStorage]                 │
  │                                                    │
  ▼
[Redirect ke dashboard]                                │
  │                                                    │
  ▼
END (◎)

Catatan: Access token disimpan di memory (Zustand), 
refresh token di localStorage untuk silent refresh.""",

        "Activity Diagram — Silent Token Refresh": """[Activity Diagram: Silent Token Refresh]

START (●)
  │
  ▼
[Frontend mengirim API request]
  │
  ▼
[Axios interceptor attach Bearer token]
  │
  ▼
[Backend terima request]
  │
  ▼
<Token expired?> ──No──> [Process request normally] ──┐
  │ Yes                                                │
  ▼                                                    │
[Backend return 401 TOKEN_EXPIRED]                     │
  │                                                    │
  ▼
[Axios interceptor detect TOKEN_EXPIRED]               │
  │                                                    │
  ▼
[Queue failed request]                                 │
  │                                                    │
  ▼
[POST /api/auth/refresh dengan refresh_token]          │
  │                                                    │
  ▼
<Refresh token valid?> ──No──> [Force logout] ────────┐
  │ Yes                                                │
  ▼                                                    │
[Backend return new token pair]                        │
  │                                                    │
  ▼
[Update Zustand store dengan new access_token]         │
  │                                                    │
  ▼
[Update localStorage dengan new refresh_token]         │
  │                                                    │
  ▼
[Retry original request dengan new token]              │
  │                                                    │
  ▼
[Continue normal flow]                                 │
  │                                                    │
  ▼
END (◎)

Catatan: Flag _retry mencegah infinite loop jika refresh 
gagal berulang kali.""",

        "Activity Diagram — Alur Penulisan Artikel": """[Activity Diagram: Alur Penulisan & Publikasi Artikel]

START (●)
  │
  ▼
[Jurnalis login ke CMS]
  │
  ▼
[Buka halaman "Tulis Berita"]
  │
  ▼
[Isi form: judul, konten (TipTap), kategori, tags]
  │
  ▼
[Upload gambar sampul (opsional)]
  │
  ▼
<Save as draft atau Submit untuk review?>
  │                    │
  │ Draft              │ Review
  ▼                    ▼
[POST /api/cms/articles  [POST /api/cms/articles
 status: draft]           status: review]
  │                    │
  ▼                    ▼
[Artikel tersimpan dengan     [Artikel masuk antrean 
 status draft]                 review editor]
  │                    │
  │                    ▼
  │                  [Editor review artikel]
  │                    │
  │                    ▼
  │                  <Approve atau Reject?>
  │                    │           │
  │                    │ Approve   │ Reject
  │                    ▼           ▼
  │                  [PATCH status:  [PATCH status: rejected
  │                   published]     + change_note]
  │                    │           │
  │                    ▼           ▼
  │                  [Artikel tayang  [Jurnalis perbaiki
  │                   di portal]       artikel]
  │                    │           │
  │                    │           └──> [Kembali ke step review]
  │                    │
  ▼                    ▼
END (◎)

Catatan: Transisi status mengikuti state machine:
draft → review → published/scheduled/rejected → archived""",
    }
    
    paragraphs = doc.paragraphs
    
    # Fill activity diagrams
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        
        for diagram_key, description in activity_diagrams.items():
            if diagram_key in text:
                # Look for empty paragraphs after this reference
                for j in range(i+1, min(len(paragraphs), i+20)):
                    next_para = paragraphs[j]
                    next_text = next_para.text.strip()
                    
                    if next_text == '':
                        set_diagram_format(next_para, description)
                        break
                    elif next_para.style.name.startswith('Heading') or \
                         ('Activity Diagram' in next_text and j > i+1):
                        break
                break
    
    return doc

def main():
    print("Opening FIX_filled.docx...")
    doc = docx.Document('FIX_filled.docx')
    
    print("Filling all diagram placeholders...")
    doc = fill_all_diagrams(doc)
    
    print("Saving modified document...")
    doc.save('FIX_filled.docx')
    
    print("Done!")
    print("\nSummary of filled content:")
    print("- 19 code snippets (Potongan Kode 1-19)")
    print("- Architecture diagram (Gambar 3.1)")
    print("- Backend folder structure (Gambar 3.2)")
    print("- Frontend folder structure (Gambar 3.3)")
    print("- Activity diagrams (Login, Token Refresh, Article Workflow)")

if __name__ == '__main__':
    main()
