from docx import Document
from PIL import Image
import io

doc = Document('laporan_pkl.docx')
NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# Map each relationship ID to its image size
rel_sizes = {}
for rel_id, rel in doc.part.rels.items():
    if 'image' in rel.reltype:
        blob = rel.target_part.blob
        try:
            img = Image.open(io.BytesIO(blob))
            rel_sizes[rel_id] = img.size
        except:
            rel_sizes[rel_id] = (0, 0)

# For each diagram title, find which image it uses and check size
import re
from lxml import etree

for i in range(455, 555):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    t = p.text.strip()
    
    # Check if this is a diagram title
    if t and ('Diagram' in t) and any(kw in t for kw in ['Activity', 'Sequence', 'Class', 'Use Case', 'State']):
        # Check next 5 paragraphs for images
        for j in range(i+1, min(i+6, len(doc.paragraphs))):
            next_p = doc.paragraphs[j]
            xml_str = etree.tostring(next_p._element, encoding='unicode')
            embeds = re.findall(r'r:embed="([^"]+)"', xml_str)
            if embeds:
                sizes = [rel_sizes.get(e, (0,0)) for e in embeds]
                min_size = min(s[0]*s[1] for s in sizes) if sizes else 0
                status = 'OK' if min_size > 10000 else f'BROKEN ({sizes})'
                print(f'[{i}] {t[:60]}')
                print(f'  -> para [{j}] embeds={embeds} sizes={sizes} [{status}]')
                break
        else:
            print(f'[{i}] {t[:60]}')
            print(f'  -> NO IMAGE FOUND in next 5 paragraphs')
