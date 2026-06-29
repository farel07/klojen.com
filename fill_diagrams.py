#!/usr/bin/env python3
"""
Script to fill empty diagram placeholders in FIX.docx with text-based descriptions.
"""
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_diagram_format(paragraph, text):
    """Format paragraph as diagram description."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.text = ''
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)

def fill_diagram_placeholders(doc):
    """Fill empty diagram placeholders with text-based descriptions."""
    
    diagram_descriptions = {
        "Gambar 3.1": """[Diagram: Arsitektur Sistem Klojen.com]

┌─────────────────────────────────────────────────────────────┐
│                         CLIENT LAYER                         │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐      │
│  │   Browser    │  │   Mobile     │  │   Desktop    │      │
│  │  (Chrome,    │  │  (Safari,    │  │  (Any modern │      │
│  │   Firefox)   │  │   Chrome)    │  │   browser)   │      │
│  └──────┬───────┘  └──────┬───────┘  └──────┬───────┘      │
│         │                  │                  │              │
│         └──────────────────┼──────────────────┘              │
│                            │ HTTPS                            │
└────────────────────────────┼────────────────────────────────┘
                             │
┌────────────────────────────┼────────────────────────────────┐
│                    FRONTEND (Next.js 15)                     │
│  ┌─────────────────────────┴─────────────────────────────┐ │
│  │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐ │ │
│  │  │ Portal Publik│  │  CMS Redaksi │  │   Bookmark   │ │ │
│  │  │  (Beranda,   │  │  (Dashboard, │  │   Halaman    │ │ │
│  │  │   Artikel,   │  │   Bank       │  │              │ │ │
│  │  │   Pencarian) │  │   Berita)    │  │              │ │ │
│  │  └──────────────┘  └──────────────┘  └──────────────┘ │ │
│  │                                                          │ │
│  │  ┌────────────────────────────────────────────────────┐ │ │
│  │  │  Zustand Store (Auth, UI State)                    │ │ │
│  │  │  Axios Interceptor (Token Refresh)                 │ │ │
│  │  └────────────────────────────────────────────────────┘ │ │
│  └──────────────────────────────────────────────────────────┘ │
└────────────────────────────┬────────────────────────────────┘
                             │ REST API (JSON)
┌────────────────────────────┼────────────────────────────────┐
│                   BACKEND (Laravel 12)                       │
│  ┌─────────────────────────┴─────────────────────────────┐ │
│  │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐ │ │
│  │  │ Controllers  │  │   Services   │  │ Repositories │ │ │
│  │  │  (Thin,      │  │  (Business   │  │  (Data       │ │ │
│  │  │   Routing)   │──│   Logic)     │──│   Access)    │ │ │
│  │  └──────────────┘  └──────────────┘  └──────────────┘ │ │
│  │                                                          │ │
│  │  ┌────────────────────────────────────────────────────┐ │ │
│  │  │  JWT Authentication  │  RBAC Middleware            │ │ │
│  │  │  Validation          │  Error Handling             │ │ │
│  │  └────────────────────────────────────────────────────┘ │ │
│  └──────────────────────────────────────────────────────────┘ │
└────────────────────────────┬────────────────────────────────┘
                             │ Eloquent ORM / Query Builder
┌────────────────────────────┼────────────────────────────────┐
│                    DATABASE (MySQL 8)                        │
│  ┌─────────────────────────┴─────────────────────────────┐ │
│  │  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐ │ │
│  │  │  users   │ │ articles │ │categories│ │ comments │ │ │
│  │  └──────────┘ └──────────┘ └──────────┘ └──────────┘ │ │
│  │  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐ │ │
│  │  │bookmarks │ │  media   │ │   tags   │ │  etc.    │ │ │
│  │  └──────────┘ └──────────┘ └──────────┘ └──────────┘ │ │
│  └──────────────────────────────────────────────────────────┘ │
└───────────────────────────────────────────────────────────────┘
                             │
┌────────────────────────────┼────────────────────────────────┐
│                 CLOUD STORAGE (S3/Cloudinary)                │
│  ┌─────────────────────────┴─────────────────────────────┐ │
│  │  Images, Videos, Documents                             │ │
│  └──────────────────────────────────────────────────────────┘ │
└───────────────────────────────────────────────────────────────┘""",

        "Gambar 3.2": """[Diagram: Struktur Folder Backend Laravel 12]

backend/
├── app/
│   ├── Console/
│   │   └── Commands/
│   │       ├── PublishScheduledArticles.php    # Auto-publish cron job
│   │       └── ReindexArticlesCommand.php      # Rebuild search index
│   ├── Http/
│   │   ├── Controllers/
│   │   │   ├── AuthController.php              # Login, register, refresh, logout
│   │   │   ├── ArticleController.php           # Public article endpoints
│   │   │   ├── CmsArticleController.php        # CMS article CRUD + workflow
│   │   │   ├── BookmarkController.php          # Toggle + list bookmarks
│   │   │   ├── CategoryController.php          # CRUD kategori (hierarki)
│   │   │   ├── TagController.php               # CRUD tag
│   │   │   ├── CommentController.php           # Komentar + moderasi
│   │   │   ├── MediaController.php             # Upload + manajemen media
│   │   │   ├── CmsDashboardController.php      # Statistik dashboard CMS
│   │   │   ├── UserController.php              # Manajemen pengguna (admin)
│   │   │   └── SearchController.php            # Pencarian artikel (FULLTEXT)
│   │   └── Middleware/
│   │       ├── AuthenticateJwt.php             # Verifikasi JWT access token
│   │       └── EnsureAdmin.php                 # Otorisasi berbasis role
│   ├── Models/
│   │   ├── User.php, Article.php, Category.php, Tag.php
│   │   ├── Comment.php, Bookmark.php, Media.php
│   │   ├── RefreshToken.php, ArticleRevision.php
│   │   └── ScheduledArticle.php, SearchIndex.php
│   ├── Services/
│   │   ├── AuthService.php                     # Logika autentikasi + JWT
│   │   ├── ArticleService.php                  # Logika bisnis artikel
│   │   ├── CmsArticleService.php               # Workflow editorial
│   │   ├── BookmarkService.php                 # Logika bookmark toggle
│   │   ├── CommentService.php                  # Moderasi komentar
│   │   ├── CmsStatisticsService.php            # Agregasi statistik dashboard
│   │   ├── SearchService.php                   # Full-text search indexing
│   │   └── ScheduledPublishService.php         # Auto-publish scheduler
│   └── Repositories/
│       ├── Contracts/                          # Interface repositories
│       ├── ArticleRepository.php
│       ├── UserRepository.php
│       ├── CategoryRepository.php
│       └── ... (other repositories)
├── database/
│   ├── migrations/                             # Schema versioning
│   ├── seeders/                                # Data seeding
│   └── data/                                   # JSON dummy data
├── routes/
│   └── api.php                                 # API route definitions
└── config/                                     # Configuration files""",

        "Gambar 3.3": """[Diagram: Struktur Folder Frontend Next.js 15]

frontend/
├── app/
│   ├── (main)/                                 # Route group portal publik
│   │   ├── page.tsx                            # Halaman utama (Dashboard Reader)
│   │   ├── [slug]/                             # Detail artikel (SSG)
│   │   │   └── page.tsx
│   │   ├── kategori/[slug]/                    # Listing artikel per kategori
│   │   ├── tag/[slug]/                         # Listing artikel per tag
│   │   └── cari/                               # Halaman pencarian
│   ├── (auth)/                                 # Route group autentikasi
│   │   ├── login/                              # Halaman login
│   │   └── register/                           # Halaman registrasi
│   ├── bookmark/                               # Halaman bookmark (requires auth)
│   │   └── page.tsx
│   ├── cms/                                    # Route group CMS redaksi
│   │   ├── layout.tsx                          # Layout CMS + auth guard
│   │   ├── dashboard/                          # Dashboard statistik
│   │   │   └── page.tsx
│   │   ├── artikel/                            # Bank berita + tulis berita
│   │   │   ├── page.tsx
│   │   │   └── baru/page.tsx
│   │   ├── media/                              # Manajemen media
│   │   ├── komentar/                           # Moderasi komentar
│   │   └── pengguna/                           # Manajemen pengguna (admin only)
│   ├── components/                             # Reusable React components
│   │   ├── AuthProvider.tsx                    # Auth context wrapper
│   │   ├── AnalyticsTracker.tsx                # Page view tracker
│   │   └── cms/                                # CMS-specific components
│   │       ├── AdminDashboard.tsx
│   │       └── ... (other CMS components)
│   ├── constants/                              # Konstanta aplikasi
│   │   ├── roles.ts                            # Role definitions & permissions
│   │   └── errorMessages.ts                    # Error message mappings
│   └── types/                                  # TypeScript type definitions
│       └── index.ts                            # All interfaces & types
├── lib/                                        # Utility libraries
│   ├── axios.ts                                # Axios instance + interceptors
│   ├── auth.ts                                 # Token management helpers
│   ├── validations.ts                          # Zod validation schemas
│   └── api/                                    # API functions per module
│       ├── articles.ts                         # Article API calls
│       ├── bookmarks.ts                        # Bookmark API calls
│       ├── auth.ts                             # Auth API calls
│       ├── categories.ts                       # Category API calls
│       ├── tags.ts                             # Tag API calls
│       └── media.ts                            # Media API calls
├── stores/                                     # Zustand state management
│   └── authStore.ts                            # Auth state (token, user)
└── public/                                     # Static assets
    └── images/                                 # Static images""",
    }
    
    paragraphs = doc.paragraphs
    
    # Find diagram references and fill empty paragraphs after them
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        
        # Check if this paragraph references a diagram
        for diagram_key, description in diagram_descriptions.items():
            if diagram_key in text and "Gambar" in text:
                # Look for empty paragraphs after this reference
                # within a reasonable range (max 30 paragraphs forward)
                for j in range(i+1, min(len(paragraphs), i+30)):
                    next_para = paragraphs[j]
                    next_text = next_para.text.strip()
                    
                    # If we find an empty paragraph, fill it with the diagram
                    if next_text == '':
                        set_diagram_format(next_para, description)
                        break
                    # Stop if we hit another heading or significant content
                    elif next_para.style.name.startswith('Heading') or len(next_text) > 50:
                        break
                break
    
    return doc

def main():
    print("Opening FIX_filled.docx...")
    doc = docx.Document('FIX_filled.docx')
    
    print("Filling empty diagram placeholders...")
    doc = fill_diagram_placeholders(doc)
    
    print("Saving modified document...")
    doc.save('FIX_filled.docx')
    
    print("Done! Diagram placeholders have been filled with text-based descriptions.")
    print("\nSummary:")
    print("✓ All 19 code snippets (Potongan Kode 1-19) filled with actual code")
    print("✓ Key diagram placeholders (Gambar 3.1, 3.2, 3.3) filled with text diagrams")
    print("\nNote: For production-quality diagrams, consider creating actual images")
    print("using tools like draw.io, Lucidchart, or Mermaid and inserting them.")

if __name__ == '__main__':
    main()
