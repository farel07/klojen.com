"""Generate diagram images and insert them + code into 23081010208_laporan pkl.docx"""
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
import docx
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
import os, io

OUT_DIR = "diagrams_laporan"
os.makedirs(OUT_DIR, exist_ok=True)

def save_fig(fig, name):
    path = os.path.join(OUT_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white', pad_inches=0.2)
    plt.close(fig)
    return path

def draw_box(ax, x, y, w, h, text, color='#E3F2FD', edge='#1565C0', fontsize=7):
    box = FancyBboxPatch((x-w/2, y-h/2), w, h, boxstyle="round,pad=0.05",
                          facecolor=color, edgecolor=edge, linewidth=1.2)
    ax.add_patch(box)
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize, fontweight='bold', wrap=True)

def draw_arrow(ax, x1, y1, x2, y2, color='#333'):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.5))

# ============================================================
# 1. ARSITEKTUR SISTEM
# ============================================================
def gen_architecture():
    fig, ax = plt.subplots(1, 1, figsize=(8, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_title('Arsitektur Decoupled / Headless CMS Klojen.com', fontsize=11, fontweight='bold', pad=10)
    
    # Client
    draw_box(ax, 5, 6.2, 4, 0.5, 'Browser (Client)', '#FFF3E0', '#E65100')
    draw_arrow(ax, 5, 5.95, 5, 5.45)
    
    # Frontend
    draw_box(ax, 5, 5.1, 6, 0.6, 'Frontend — Next.js 15\n(SSR / SSG / CSR + TypeScript + Tailwind CSS)', '#E8F5E9', '#2E7D32')
    draw_arrow(ax, 5, 4.8, 5, 4.3)
    ax.text(5.8, 4.55, 'REST API (JSON)\nBearer JWT', fontsize=6, style='italic', color='#555')
    
    # Middleware
    draw_box(ax, 5, 3.9, 5, 0.5, 'Laravel 12 Middleware (JWT Auth + RBAC + Rate Limit)', '#FFF9C4', '#F57F17')
    draw_arrow(ax, 5, 3.65, 5, 3.15)
    
    # Controller
    draw_box(ax, 5, 2.8, 4, 0.5, 'Controllers (Thin)', '#E3F2FD', '#1565C0')
    draw_arrow(ax, 5, 2.55, 5, 2.05)
    
    # Service
    draw_box(ax, 5, 1.7, 4, 0.5, 'Services (Business Logic)', '#F3E5F5', '#7B1FA2')
    draw_arrow(ax, 5, 1.45, 5, 0.95)
    
    # Repository
    draw_box(ax, 5, 0.6, 4, 0.5, 'Repositories (Data Access)', '#FFEBEE', '#C62828')
    draw_arrow(ax, 5, 0.35, 5, -0.15)
    
    # Database
    draw_box(ax, 5, -0.5, 3.5, 0.5, 'MySQL 8 Database\n(14 Tabel + FULLTEXT Index)', '#ECEFF1', '#37474F')
    
    # Side labels
    ax.text(0.3, 5.1, 'Frontend\n(Portal + CMS)', fontsize=6, color='#2E7D32', fontweight='bold', va='center')
    ax.text(8.5, 3.3, 'Backend\n(Laravel 12\nAPI)', fontsize=6, color='#1565C0', fontweight='bold', va='center')
    
    return save_fig(fig, 'arsitektur.png')

# ============================================================
# 2. SERVICE-REPOSITORY PATTERN
# ============================================================
def gen_service_repo():
    fig, ax = plt.subplots(1, 1, figsize=(7, 3.5))
    ax.set_xlim(0, 9); ax.set_ylim(0, 5)
    ax.axis('off')
    ax.set_title('Service-Repository Pattern', fontsize=11, fontweight='bold')
    
    draw_box(ax, 1.5, 4, 2, 0.7, 'HTTP\nRequest', '#FFF3E0', '#E65100')
    draw_arrow(ax, 2.5, 4, 3.5, 4)
    draw_box(ax, 4.5, 4, 2, 0.7, 'Controller\n(thin)', '#E3F2FD', '#1565C0')
    draw_arrow(ax, 5.5, 4, 6.5, 4)
    draw_box(ax, 7.5, 4, 2, 0.7, 'Service\n(logic)', '#F3E5F5', '#7B1FA2')
    
    draw_arrow(ax, 7.5, 3.65, 4.5, 2.35)
    draw_box(ax, 4.5, 2, 3, 0.7, 'Repository\n(interface)', '#FFEBEE', '#C62828')
    draw_arrow(ax, 4.5, 1.65, 4.5, 0.85)
    draw_box(ax, 4.5, 0.5, 3, 0.7, 'Database\n(MySQL 8)', '#ECEFF1', '#37474F')
    
    # Labels
    ax.text(1, 3.2, 'delegate', fontsize=6, style='italic', color='#555')
    ax.text(6, 3.2, 'call interface', fontsize=6, style='italic', color='#555')
    ax.text(5.5, 1.2, 'Eloquent / QB', fontsize=6, style='italic', color='#555')
    
    return save_fig(fig, 'service_repo.png')

# ============================================================
# 3. FOLDER STRUCTURE (as image)
# ============================================================
def gen_backend_folder():
    fig, ax = plt.subplots(1, 1, figsize=(6, 5))
    ax.axis('off')
    text = """backend/ (Laravel 12)
├── app/
│   ├── Console/Commands/
│   │   ├── PublishScheduledArticles.php
│   │   └── ReindexArticlesCommand.php
│   ├── Http/
│   │   ├── Controllers/ (16 files)
│   │   └── Middleware/EnsureAdmin.php
│   ├── Models/ (12 models)
│   ├── Providers/
│   │   └── RepositoryServiceProvider.php
│   ├── Repositories/
│   │   ├── Contracts/ (9 interfaces)
│   │   └── (7 implementations)
│   └── Services/ (12 services)
├── database/
│   ├── migrations/ (26 files)
│   ├── seeders/ (8 files)
│   └── data/dummy2.json
├── routes/api.php (132 lines)
└── config/ (jwt, cors, db, mail...)"""
    ax.text(0.05, 0.95, text, fontsize=7, fontfamily='monospace', va='top',
            transform=ax.transAxes, bbox=dict(boxstyle='round', facecolor='#F5F5F5', edgecolor='#999'))
    ax.set_title('Struktur Folder Backend', fontsize=10, fontweight='bold')
    return save_fig(fig, 'backend_folder.png')

def gen_frontend_folder():
    fig, ax = plt.subplots(1, 1, figsize=(6, 5))
    ax.axis('off')
    text = """frontend/ (Next.js 15)
├── app/
│   ├── (main)/           # Portal Publik
│   │   ├── page.tsx      # Dashboard Reader
│   │   ├── [slug]/       # Detail Artikel (SSG)
│   │   └── bookmark/
│   ├── cms/              # CMS Redaksi
│   │   ├── dashboard/
│   │   ├── artikel/
│   │   │   ├── tulis/
│   │   │   └── tinjauan/
│   │   ├── media/
│   │   ├── komentar/
│   │   └── pengguna/
│   ├── components/
│   ├── constants/ (roles.ts, errorMessages.ts)
│   ├── types/index.ts
│   └── layout.tsx
├── lib/
│   ├── api/ (articles, bookmarks, ...)
│   ├── axios.ts (interceptors)
│   └── validations.ts (Zod)
├── stores/authStore.ts (Zustand)
└── utils/cropImage.ts"""
    ax.text(0.05, 0.95, text, fontsize=7, fontfamily='monospace', va='top',
            transform=ax.transAxes, bbox=dict(boxstyle='round', facecolor='#F5F5F5', edgecolor='#999'))
    ax.set_title('Struktur Folder Frontend', fontsize=10, fontweight='bold')
    return save_fig(fig, 'frontend_folder.png')

# ============================================================
# 4. USE CASE DIAGRAM
# ============================================================
def gen_usecase():
    fig, ax = plt.subplots(1, 1, figsize=(8, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title('Use Case Diagram — Portal Berita Klojen.com', fontsize=11, fontweight='bold')
    
    # Actors
    actors = [('Pengunjung', 0.8, 7), ('Reader', 0.8, 5.8), ('Jurnalis', 0.8, 4.2),
              ('Editor', 0.8, 2.8), ('Admin', 0.8, 1.2)]
    for name, x, y in actors:
        ax.plot(x, y+0.25, 'o', markersize=8, color='#1565C0')
        ax.plot([x, x], [y+0.15, y-0.15], '-', color='#1565C0', lw=2)
        ax.plot([x-0.15, x+0.15], [y+0.05, y+0.05], '-', color='#1565C0', lw=2)
        ax.plot([x-0.12, x], [y-0.15, y-0.3], '-', color='#1565C0', lw=2)
        ax.plot([x+0.12, x], [y-0.15, y-0.3], '-', color='#1565C0', lw=2)
        ax.text(x, y-0.5, name, ha='center', fontsize=7, fontweight='bold')
    
    # System boundary
    rect = plt.Rectangle((2.5, 0.3), 7, 7.2, fill=False, edgecolor='#333', linewidth=1.5, linestyle='--')
    ax.add_patch(rect)
    ax.text(6, 7.7, 'Sistem Portal Berita Klojen.com', ha='center', fontsize=9, fontweight='bold')
    
    # Use cases
    ucs = [
        (4.5, 7, 'Lihat Beranda'), (7, 7, 'Cari Artikel'),
        (4.5, 6.2, 'Lihat Detail Artikel'), (7, 6.2, 'Bookmark Artikel'),
        (4.5, 5.4, 'Komentar Artikel'), (7, 5.4, 'Login/Logout'),
        (4.5, 4.2, 'Tulis Artikel'), (7, 4.2, 'Edit Artikel'),
        (4.5, 3.4, 'Submit untuk Review'), (7, 3.4, 'Upload Media'),
        (4.5, 2.5, 'Review Artikel'), (7, 2.5, 'Publish/Reject'),
        (4.5, 1.7, 'Kelola Komentar'), (7, 1.7, 'Schedule Publish'),
        (4.5, 0.8, 'Kelola Pengguna'), (7, 0.8, 'Lihat Dashboard'),
    ]
    for x, y, label in ucs:
        ellipse = mpatches.Ellipse((x, y), 2.2, 0.55, facecolor='#E3F2FD', edgecolor='#1565C0', lw=1)
        ax.add_patch(ellipse)
        ax.text(x, y, label, ha='center', va='center', fontsize=6.5)
    
    # Lines (simplified connections)
    connections = [
        (0.8, 7, 3.4, 7), (0.8, 7, 3.4, 6.2),  # Pengunjung
        (0.8, 5.8, 3.4, 5.4), (0.8, 5.8, 3.4, 6.2), (0.8, 5.8, 5.9, 5.4),
        (0.8, 4.2, 3.4, 4.2), (0.8, 4.2, 3.4, 3.4), (0.8, 4.2, 5.9, 4.2),
        (0.8, 2.8, 3.4, 2.5), (0.8, 2.8, 3.4, 3.4), (0.8, 2.8, 5.9, 2.5),
        (0.8, 1.2, 3.4, 0.8), (0.8, 1.2, 3.4, 2.5), (0.8, 1.2, 5.9, 0.8),
    ]
    for x1, y1, x2, y2 in connections:
        ax.plot([x1+0.2, x2], [y1, y2], '-', color='#999', lw=0.5, alpha=0.5)
    
    return save_fig(fig, 'usecase.png')

# ============================================================
# 5. ACTIVITY DIAGRAMS (horizontal layout per memory)
# ============================================================
def gen_activity_diagram(title, steps, filename):
    n = len(steps)
    fig, ax = plt.subplots(1, 1, figsize=(max(8, n*1.3), 2.2))
    ax.set_xlim(-0.5, n*1.3+0.5); ax.set_ylim(-0.5, 2.5)
    ax.axis('off')
    ax.set_title(title, fontsize=10, fontweight='bold')
    
    # Start node
    circle = plt.Circle((0, 1), 0.2, color='#333')
    ax.add_patch(circle)
    
    for i, (step, is_decision) in enumerate(steps):
        x = (i+1) * 1.3
        if is_decision == 'diamond':
            diamond = plt.Polygon([(x, 1.4), (x+0.5, 1), (x, 0.6), (x-0.5, 1)],
                                   facecolor='#FFF9C4', edgecolor='#F57F17', lw=1.2)
            ax.add_patch(diamond)
            ax.text(x, 1, step, ha='center', va='center', fontsize=5.5, fontweight='bold')
        elif is_decision == 'end':
            circle = plt.Circle((x, 1), 0.2, color='#333')
            ax.add_patch(circle)
            inner = plt.Circle((x, 1), 0.12, color='white')
            ax.add_patch(inner)
        else:
            box = FancyBboxPatch((x-0.5, 0.6), 1, 0.8, boxstyle="round,pad=0.05",
                                  facecolor='#E3F2FD', edgecolor='#1565C0', lw=1)
            ax.add_patch(box)
            ax.text(x, 1, step, ha='center', va='center', fontsize=5.5, wrap=True)
        
        # Arrow
        ax.annotate('', xy=(x-0.55 if is_decision != 'end' else x-0.25, 1),
                    xytext=(x-0.55 if i == 0 else (i)*1.3+0.55, 1),
                    arrowprops=dict(arrowstyle='->', color='#333', lw=1.2))
    
    return save_fig(fig, filename)

def gen_all_activities():
    paths = {}
    
    # Login
    paths['login'] = gen_activity_diagram('Activity Diagram — Login & Autentikasi', [
        ('Buka halaman\nlogin', 'box'), ('Input email\n& password', 'box'),
        ('Validasi\ninput', 'diamond'), ('POST /auth/login', 'box'),
        ('Verifikasi\nkredensial', 'diamond'), ('Generate JWT\n+ Refresh Token', 'box'),
        ('Simpan token\nke store', 'box'), ('Redirect ke\ndashboard', 'box'), ('', 'end'),
    ], 'act_login.png')
    
    # Token Refresh
    paths['refresh'] = gen_activity_diagram('Activity Diagram — Silent Token Refresh', [
        ('Request ke\nAPI', 'box'), ('Token\nexpired?', 'diamond'),
        ('Interceptor\ndeteksi error', 'box'), ('POST /auth/refresh', 'box'),
        ('Verifikasi\nrefresh token', 'diamond'), ('Generate access\ntoken baru', 'box'),
        ('Update store\n& retry request', 'box'), ('Response\nberhasil', 'box'), ('', 'end'),
    ], 'act_refresh.png')
    
    # Article Workflow
    paths['article'] = gen_activity_diagram('Activity Diagram — Alur Penulisan & Publikasi Artikel', [
        ('Jurnalis buat\nartikel baru', 'box'), ('Status =\ndraft', 'box'),
        ('Edit & lengkapi\ndata', 'box'), ('Submit untuk\nreview', 'box'),
        ('Editor\nreview', 'diamond'), ('Setujui →\npublished', 'box'),
        ('Tolak →\nrejected', 'box'), ('Artikel tayang\ndi portal', 'box'), ('', 'end'),
    ], 'act_article.png')
    
    # Comment Moderation
    paths['comment'] = gen_activity_diagram('Activity Diagram — Moderasi Komentar', [
        ('Pembaca tulis\nkomentar', 'box'), ('Validasi input\n& rate limit', 'diamond'),
        ('POST comment\nke API', 'box'), ('Cek parent_id\n(threaded?)', 'diamond'),
        ('Simpan ke DB\nstatus=pending', 'box'), ('Editor review\nkomentar', 'box'),
        ('Approve /\nReject', 'diamond'), ('Update status\ndi database', 'box'), ('', 'end'),
    ], 'act_comment.png')
    
    # Scheduled Publishing
    paths['scheduled'] = gen_activity_diagram('Activity Diagram — Scheduled Publishing', [
        ('Editor set\nstatus=scheduled', 'box'), ('Input\nscheduled_at', 'box'),
        ('Validasi min\n5 menit', 'diamond'), ('Simpan ke\nscheduled_articles', 'box'),
        ('Cron job\nsetiap menit', 'box'), ('Cek scheduled_at\n<= NOW()', 'diamond'),
        ('Update status\n= published', 'box'), ('Reindex\nsearch_indexes', 'box'), ('', 'end'),
    ], 'act_scheduled.png')
    
    # Upload Media
    paths['upload'] = gen_activity_diagram('Activity Diagram — Upload Media', [
        ('Pilih file\ndari device', 'box'), ('Validasi client\ntipe & ukuran', 'diamond'),
        ('Kirim file\nmultipart', 'box'), ('Validasi server\ntipe & ukuran', 'diamond'),
        ('Upload ke\ncloud storage', 'box'), ('Simpan metadata\nke database', 'box'),
        ('Return URL\nmedia', 'box'), ('', 'end'),
    ], 'act_upload.png')
    
    # Bookmark
    paths['bookmark'] = gen_activity_diagram('Activity Diagram — Bookmark Artikel (Toggle)', [
        ('Klik tombol\nbookmark', 'box'), ('Cek status\nbookmark', 'diamond'),
        ('Sudah ada?', 'diamond'), ('Hapus dari\ndatabase', 'box'),
        ('Update UI:\nunbookmark', 'box'), ('Insert ke\ndatabase', 'box'),
        ('Update UI:\nbookmarked', 'box'), ('Response ke\nfrontend', 'box'), ('', 'end'),
    ], 'act_bookmark.png')
    
    # Search
    paths['search'] = gen_activity_diagram('Activity Diagram — Pencarian Artikel', [
        ('Input kata\nkunci', 'box'), ('Validasi\ninput', 'diamond'),
        ('GET /search?q=keyword', 'box'), ('FULLTEXT query\nMySQL', 'box'),
        ('MATCH AGAINST\nnatural mode', 'box'), ('Ranking by\nrelevansi', 'box'),
        ('Return hasil\n+ pagination', 'box'), ('', 'end'),
    ], 'act_search.png')
    
    # Logout
    paths['logout'] = gen_activity_diagram('Activity Diagram — Logout & Invalidasi Sesi', [
        ('Klik tombol\nlogout', 'box'), ('POST /auth/logout', 'box'),
        ('Revoke refresh\ntoken di DB', 'box'), ('Hapus token\ndari store', 'box'),
        ('Clear localStorage', 'box'), ('Reset Zustand\nstate', 'box'),
        ('Redirect ke\n/login', 'box'), ('', 'end'),
    ], 'act_logout.png')
    
    return paths

# ============================================================
# 6. SEQUENCE DIAGRAMS
# ============================================================
def gen_sequence_diagram(title, participants, messages, filename):
    n = len(participants)
    fig, ax = plt.subplots(1, 1, figsize=(max(7, n*1.8), max(4, len(messages)*0.45+1.5)))
    ax.set_xlim(-0.5, n*1.8+0.5); ax.set_ylim(-0.5, len(messages)*0.45+2)
    ax.axis('off')
    ax.set_title(title, fontsize=10, fontweight='bold')
    
    # Participant boxes
    for i, p in enumerate(participants):
        x = i * 1.8 + 0.5
        box = FancyBboxPatch((x-0.6, len(messages)*0.45+0.8), 1.2, 0.5,
                              facecolor='#1a3a6b', edgecolor='#0d2240', lw=1.5)
        ax.add_patch(box)
        ax.text(x, len(messages)*0.45+1.05, p, ha='center', va='center',
                fontsize=7, color='white', fontweight='bold')
        # Lifeline
        ax.plot([x, x], [0.3, len(messages)*0.45+0.8], '--', color='#aaa', lw=0.8)
    
    # Messages
    for i, (src, dst, label, is_return) in enumerate(messages):
        y = len(messages)*0.45 + 0.5 - i*0.45
        src_idx = participants.index(src) if src in participants else 0
        dst_idx = participants.index(dst) if dst in participants else 1
        x1 = src_idx * 1.8 + 0.5
        x2 = dst_idx * 1.8 + 0.5
        
        style = '-' if is_return else '->'
        color = '#666' if is_return else '#222'
        ls = 'dashed' if is_return else 'solid'
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle=style, color=color, lw=1.2, linestyle=ls))
        mid_x = (x1+x2)/2
        ax.text(mid_x, y+0.12, label, ha='center', fontsize=5.5, color='#333')
    
    return save_fig(fig, filename)

def gen_all_sequences():
    paths = {}
    
    paths['seq_login'] = gen_sequence_diagram('Sequence Diagram — Login',
        ['User', 'Frontend', 'API', 'AuthService', 'DB'],
        [('User', 'Frontend', 'Input email+password', False),
         ('Frontend', 'API', 'POST /auth/login', False),
         ('API', 'AuthService', 'verify credentials', False),
         ('AuthService', 'DB', 'Query user', False),
         ('DB', 'AuthService', 'User data', True),
         ('AuthService', 'DB', 'Store refresh token', False),
         ('AuthService', 'API', 'JWT tokens', True),
         ('API', 'Frontend', 'access_token + refresh_token', True),
         ('Frontend', 'Frontend', 'Save to Zustand store', False)],
        'seq_login.png')
    
    paths['seq_refresh'] = gen_sequence_diagram('Sequence Diagram — Silent Token Refresh',
        ['Frontend', 'Interceptor', 'API', 'AuthService', 'DB'],
        [('Frontend', 'Interceptor', 'Request with expired token', False),
         ('Interceptor', 'API', 'Forward request', False),
         ('API', 'Interceptor', 'TOKEN_EXPIRED error', True),
         ('Interceptor', 'API', 'POST /auth/refresh', False),
         ('API', 'AuthService', 'Validate refresh token', False),
         ('AuthService', 'DB', 'Check refresh_tokens', False),
         ('DB', 'AuthService', 'Token valid', True),
         ('AuthService', 'API', 'New access_token', True),
         ('API', 'Interceptor', 'New token', True),
         ('Interceptor', 'API', 'Retry original request', False)],
        'seq_refresh.png')
    
    paths['seq_crud'] = gen_sequence_diagram('Sequence Diagram — CRUD Artikel CMS',
        ['User', 'Frontend', 'API', 'CmsArticleService', 'DB'],
        [('User', 'Frontend', 'Aksi CRUD artikel', False),
         ('Frontend', 'API', 'HTTP Request + JWT', False),
         ('API', 'CmsArticleService', 'Delegate operation', False),
         ('CmsArticleService', 'DB', 'Query/Insert/Update/Delete', False),
         ('DB', 'CmsArticleService', 'Result', True),
         ('CmsArticleService', 'API', 'Response data', True),
         ('API', 'Frontend', 'JSON response', True),
         ('Frontend', 'User', 'Update UI', False)],
        'seq_crud.png')
    
    paths['seq_review'] = gen_sequence_diagram('Sequence Diagram — Review & Publish Artikel',
        ['Editor', 'Frontend', 'API', 'CmsArticleService', 'DB'],
        [('Editor', 'Frontend', 'Buka halaman tinjauan', False),
         ('Frontend', 'API', 'GET /cms/articles?status=review', False),
         ('API', 'CmsArticleService', 'Get review articles', False),
         ('CmsArticleService', 'DB', 'Query articles', False),
         ('DB', 'CmsArticleService', 'Article list', True),
         ('Editor', 'Frontend', 'Klik Publish/Reject', False),
         ('Frontend', 'API', 'PATCH /articles/{id}/status', False),
         ('API', 'CmsArticleService', 'Validate transition', False),
         ('CmsArticleService', 'DB', 'Update status', False),
         ('CmsArticleService', 'API', 'Updated article', True)],
        'seq_review.png')
    
    paths['seq_bookmark'] = gen_sequence_diagram('Sequence Diagram — Bookmark Artikel',
        ['User', 'Frontend', 'API', 'BookmarkService', 'DB'],
        [('User', 'Frontend', 'Klik tombol bookmark', False),
         ('Frontend', 'API', 'POST /bookmarks {article_id}', False),
         ('API', 'BookmarkService', 'toggle(userId, articleId)', False),
         ('BookmarkService', 'DB', 'Check existing bookmark', False),
         ('DB', 'BookmarkService', 'Result', True),
         ('BookmarkService', 'DB', 'Insert/Delete', False),
         ('BookmarkService', 'API', 'bookmarked: bool', True),
         ('API', 'Frontend', 'JSON response', True),
         ('Frontend', 'User', 'Toggle UI state', False)],
        'seq_bookmark.png')
    
    paths['seq_analytics'] = gen_sequence_diagram('Sequence Diagram — Page View Analytics',
        ['Browser', 'Tracker', 'API', 'AnalyticsController', 'DB'],
        [('Browser', 'Tracker', 'Navigasi halaman', False),
         ('Tracker', 'API', 'POST /analytics/track', False),
         ('API', 'AnalyticsController', 'Track page view', False),
         ('AnalyticsController', 'DB', 'Check dedup (5 min)', False),
         ('DB', 'AnalyticsController', 'Last view time', True),
         ('AnalyticsController', 'DB', 'Insert page_view', False),
         ('AnalyticsController', 'API', 'Success', True)],
        'seq_analytics.png')
    
    paths['seq_register'] = gen_sequence_diagram('Sequence Diagram — Registrasi Akun Baru',
        ['User', 'Frontend', 'API', 'AuthService', 'DB'],
        [('User', 'Frontend', 'Input data registrasi', False),
         ('Frontend', 'Frontend', 'Validasi Zod schema', False),
         ('Frontend', 'API', 'POST /auth/register', False),
         ('API', 'AuthService', 'Register user', False),
         ('AuthService', 'DB', 'Check email unique', False),
         ('DB', 'AuthService', 'Available', True),
         ('AuthService', 'DB', 'Insert user (bcrypt)', False),
         ('AuthService', 'DB', 'Send email credentials', False),
         ('AuthService', 'API', 'User created', True)],
        'seq_register.png')
    
    paths['seq_list'] = gen_sequence_diagram('Sequence Diagram — Ambil Daftar Artikel',
        ['User', 'Frontend', 'API', 'ArticleRepository', 'DB'],
        [('User', 'Frontend', 'Buka halaman beranda', False),
         ('Frontend', 'API', 'GET /articles?page=1', False),
         ('API', 'ArticleRepository', 'FindAll with pagination', False),
         ('ArticleRepository', 'DB', 'SELECT + JOIN + LIMIT', False),
         ('DB', 'ArticleRepository', 'Articles + relations', True),
         ('ArticleRepository', 'API', 'Article data', True),
         ('API', 'Frontend', 'JSON + pagination', True),
         ('Frontend', 'User', 'Render article cards', False)],
        'seq_list.png')
    
    paths['seq_submit'] = gen_sequence_diagram('Sequence Diagram — Submit Artikel Baru (Jurnalis)',
        ['Jurnalis', 'Frontend', 'API', 'CmsArticleService', 'DB'],
        [('Jurnalis', 'Frontend', 'Isi form artikel', False),
         ('Frontend', 'Frontend', 'Validasi Zod schema', False),
         ('Frontend', 'API', 'POST /cms/articles (multipart)', False),
         ('API', 'CmsArticleService', 'Create article', False),
         ('CmsArticleService', 'DB', 'Insert article', False),
         ('CmsArticleService', 'DB', 'Insert article_tags', False),
         ('CmsArticleService', 'API', 'Created article', True),
         ('API', 'Frontend', 'JSON response', True)],
        'seq_submit.png')
    
    paths['seq_comment'] = gen_sequence_diagram('Sequence Diagram — Post Komentar',
        ['Reader', 'Frontend', 'API', 'CommentService', 'DB'],
        [('Reader', 'Frontend', 'Isi form komentar', False),
         ('Frontend', 'Frontend', 'Validasi + rate limit check', False),
         ('Frontend', 'API', 'POST /articles/{id}/comments', False),
         ('API', 'CommentService', 'Create comment', False),
         ('CommentService', 'DB', 'Check rate limit', False),
         ('CommentService', 'DB', 'Insert comment', False),
         ('CommentService', 'API', 'Created comment', True),
         ('API', 'Frontend', 'JSON response', True)],
        'seq_comment.png')
    
    paths['seq_upload'] = gen_sequence_diagram('Sequence Diagram — Upload Media',
        ['User', 'Frontend', 'API', 'MediaService', 'Cloud', 'DB'],
        [('User', 'Frontend', 'Pilih file media', False),
         ('Frontend', 'Frontend', 'Validasi client-side', False),
         ('Frontend', 'API', 'POST /media (multipart)', False),
         ('API', 'MediaService', 'Upload media', False),
         ('MediaService', 'MediaService', 'Validasi server-side', False),
         ('MediaService', 'Cloud', 'Upload file', False),
         ('Cloud', 'MediaService', 'File URL', True),
         ('MediaService', 'DB', 'Save metadata', False),
         ('MediaService', 'API', 'Media data + URL', True)],
        'seq_upload.png')
    
    paths['seq_edit'] = gen_sequence_diagram('Sequence Diagram — Edit & Update Artikel',
        ['Jurnalis', 'Frontend', 'API', 'CmsArticleService', 'DB'],
        [('Jurnalis', 'Frontend', 'Buka form edit', False),
         ('Frontend', 'API', 'GET /cms/articles/{id}', False),
         ('API', 'CmsArticleService', 'Get article', False),
         ('Jurnalis', 'Frontend', 'Edit data artikel', False),
         ('Frontend', 'API', 'PUT /cms/articles/{id}', False),
         ('API', 'CmsArticleService', 'Validate ownership', False),
         ('CmsArticleService', 'DB', 'Save revision (audit)', False),
         ('CmsArticleService', 'DB', 'Update article', False),
         ('CmsArticleService', 'API', 'Updated article', True)],
        'seq_edit.png')
    
    return paths

# ============================================================
# 7. CLASS DIAGRAM
# ============================================================
def gen_class_backend():
    fig, ax = plt.subplots(1, 1, figsize=(8, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_title('Class Diagram — Backend (Service-Repository Pattern)', fontsize=10, fontweight='bold')
    
    classes = [
        (1.5, 6, 'AuthController', ['+login()', '+register()', '+logout()', '+refresh()']),
        (4, 6, 'CmsArticleController', ['+index()', '+store()', '+update()', '+destroy()']),
        (7, 6, 'BookmarkController', ['+index()', '+toggle()']),
        (1.5, 3.5, 'AuthService', ['+login(): TokenPair', '+register(): User', '+refresh(): Token']),
        (4, 3.5, 'CmsArticleService', ['+create(): Article', '+updateStatus(): Article', '+validateTransition()']),
        (7, 3.5, 'BookmarkService', ['+toggle(): bool', '+getBookmarks(): array']),
        (1.5, 1, 'ArticleRepository', ['+findAll(): Collection', '+findBySlug(): Article', '+create(): Article']),
        (4, 1, 'UserRepository', ['+findById(): User', '+findByEmail(): User']),
        (7, 1, 'BookmarkRepository', ['+findByUser(): Collection', '+toggle(): bool']),
    ]
    
    for x, y, name, methods in classes:
        box = FancyBboxPatch((x-0.9, y-0.7), 1.8, 1.4, boxstyle="round,pad=0.05",
                              facecolor='#E3F2FD', edgecolor='#1565C0', lw=1.2)
        ax.add_patch(box)
        ax.text(x, y+0.35, name, ha='center', fontsize=6, fontweight='bold')
        ax.plot([x-0.85, x+0.85], [y+0.15, y+0.15], '-', color='#1565C0', lw=0.5)
        for i, m in enumerate(methods[:3]):
            ax.text(x, y-0.05-i*0.2, m, ha='center', fontsize=4.5, fontfamily='monospace')
    
    # Arrows (controller -> service -> repository)
    for cx, sx, rx in [(1.5, 1.5, 1.5), (4, 4, 4), (7, 7, 7)]:
        ax.annotate('', xy=(cx, 4.2), xytext=(cx, 5.3),
                    arrowprops=dict(arrowstyle='->', color='#1565C0', lw=1, linestyle='dashed'))
        ax.annotate('', xy=(rx, 1.7), xytext=(rx, 2.8),
                    arrowprops=dict(arrowstyle='->', color='#C62828', lw=1, linestyle='dashed'))
    
    return save_fig(fig, 'class_backend.png')

def gen_class_frontend():
    fig, ax = plt.subplots(1, 1, figsize=(7, 4))
    ax.set_xlim(0, 9); ax.set_ylim(0, 5)
    ax.axis('off')
    ax.set_title('Class Diagram — Frontend Architecture', fontsize=10, fontweight='bold')
    
    comps = [
        (1.5, 4, 'AuthStore\n(Zustand)', ['# accessToken', '# user', '+ setAuth()', '+ logout()']),
        (4.5, 4, 'AxiosInstance', ['# interceptors', '+ request()', '+ response()']),
        (7.5, 4, 'Zod Schemas', ['# loginSchema', '# articleSchema', '# commentSchema']),
        (1.5, 1.5, 'DashboardPage', ['+ role-based render', '+ stat cards']),
        (4.5, 1.5, 'ArticleList', ['+ pagination', '+ filter', '+ search']),
        (7.5, 1.5, 'BookmarkPage', ['+ grid cards', '+ toggle', '+ pagination']),
    ]
    
    for x, y, name, details in comps:
        box = FancyBboxPatch((x-1, y-0.7), 2, 1.4, boxstyle="round,pad=0.05",
                              facecolor='#E8F5E9', edgecolor='#2E7D32', lw=1.2)
        ax.add_patch(box)
        ax.text(x, y+0.3, name, ha='center', fontsize=6.5, fontweight='bold')
        for i, d in enumerate(details[:3]):
            ax.text(x, y-0.05-i*0.2, d, ha='center', fontsize=4.5, fontfamily='monospace')
    
    # Dependencies
    ax.annotate('', xy=(2.5, 3.3), xytext=(3.5, 3.3),
                arrowprops=dict(arrowstyle='->', color='#666', lw=1, linestyle='dashed'))
    ax.annotate('', xy=(5.5, 3.3), xytext=(6.5, 3.3),
                arrowprops=dict(arrowstyle='->', color='#666', lw=1, linestyle='dashed'))
    
    return save_fig(fig, 'class_frontend.png')

# ============================================================
# 8. CDM (Conceptual Data Model)
# ============================================================
def gen_cdm():
    fig, ax = plt.subplots(1, 1, figsize=(8, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title('Conceptual Data Model (CDM)', fontsize=10, fontweight='bold')
    
    entities = {
        'users': (5, 4), 'articles': (2, 4), 'categories': (2, 6.5),
        'tags': (8, 6.5), 'comments': (8, 4), 'bookmarks': (2, 1.5),
        'media': (5, 1.5), 'article_tags': (5, 6.5),
        'refresh_tokens': (8, 1.5), 'search_indexes': (2, 6.5),
    }
    
    main_entities = [
        (2, 4, 'USERS', '#E3F2FD'), (5, 4, 'ARTICLES', '#FFF9C4'),
        (2, 6.5, 'CATEGORIES', '#E8F5E9'), (8, 6.5, 'TAGS', '#F3E5F5'),
        (8, 4, 'COMMENTS', '#FFEBEE'), (2, 1.5, 'BOOKMARKS', '#ECEFF1'),
        (5, 1.5, 'MEDIA', '#FFF3E0'), (5, 6.5, 'ARTICLE_TAGS', '#E0F7FA'),
        (8, 1.5, 'REFRESH_TOKENS', '#FBE9E7'),
    ]
    
    for x, y, name, color in main_entities:
        box = FancyBboxPatch((x-0.7, y-0.35), 1.4, 0.7, boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor='#333', lw=1.2)
        ax.add_patch(box)
        ax.text(x, y, name, ha='center', va='center', fontsize=6.5, fontweight='bold')
    
    # Relationships
    rels = [
        (2, 4, 5, 4, 'menulis\n(1:N)'), (2, 6.5, 5, 4, 'kategori\n(1:N)'),
        (5, 4, 8, 6.5, 'tag\n(M:N)'), (5, 4, 8, 4, 'komentar\n(1:N)'),
        (5, 4, 2, 1.5, 'bookmark\n(1:N)'), (5, 4, 5, 6.5, 'pivot\n(M:N)'),
        (5, 4, 5, 1.5, 'media\n(1:N)'), (2, 4, 8, 1.5, 'token\n(1:N)'),
    ]
    for x1, y1, x2, y2, label in rels:
        ax.plot([x1, x2], [y1, y2], '-', color='#999', lw=0.8)
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx, my+0.15, label, ha='center', fontsize=5, color='#555', style='italic')
    
    return save_fig(fig, 'cdm.png')

# ============================================================
# 9. EDITORIAL WORKFLOW STATE MACHINE
# ============================================================
def gen_editorial_workflow():
    fig, ax = plt.subplots(1, 1, figsize=(8, 3))
    ax.set_xlim(-0.5, 10); ax.set_ylim(0, 3.5)
    ax.axis('off')
    ax.set_title('Alur Kerja Editorial (State Machine)', fontsize=10, fontweight='bold')
    
    states = [
        (0.5, 2, 'Draft', '#E3F2FD'), (2.5, 2, 'Review', '#FFF9C4'),
        (5, 2, 'Published', '#E8F5E9'), (7.5, 2, 'Archived', '#ECEFF1'),
        (5, 0.5, 'Rejected', '#FFEBEE'), (2.5, 0.5, 'Scheduled', '#F3E5F5'),
    ]
    
    for x, y, name, color in states:
        box = FancyBboxPatch((x-0.6, y-0.35), 1.2, 0.7, boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor='#333', lw=1.5)
        ax.add_patch(box)
        ax.text(x, y, name, ha='center', va='center', fontsize=7, fontweight='bold')
    
    # Transitions
    transitions = [
        (1.1, 2, 1.9, 2, 'submit'), (3.1, 2, 4.4, 2, 'approve'),
        (5.6, 2, 6.9, 2, 'archive'), (5, 1.65, 5, 0.85, 'reject'),
        (4.4, 0.5, 1.1, 1.65, 'revise'), (2.5, 1.65, 2.5, 0.85, 'schedule'),
    ]
    for x1, y1, x2, y2, label in transitions:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#1565C0', lw=1.2))
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx+0.1, my+0.1, label, fontsize=5.5, color='#1565C0', style='italic')
    
    return save_fig(fig, 'editorial_workflow.png')

# ============================================================
# GENERATE ALL DIAGRAMS
# ============================================================
print("Generating diagrams...")
img_arch = gen_architecture()
img_sr = gen_service_repo()
img_bf = gen_backend_folder()
img_ff = gen_frontend_folder()
img_uc = gen_usecase()
act_imgs = gen_all_activities()
seq_imgs = gen_all_sequences()
img_cb = gen_class_backend()
img_cf = gen_class_frontend()
img_cdm = gen_cdm()
img_ew = gen_editorial_workflow()
print(f"Generated {2 + len(act_imgs) + len(seq_imgs) + 4} diagram images")

# ============================================================
# NOW INSERT INTO DOCX
# ============================================================
print("\nInserting into document...")
doc = docx.Document('23081010208_laporan pkl.docx')
paragraphs = doc.paragraphs

def insert_image_after(paragraph, img_path, width_inches=5.5):
    """Insert an image as a new paragraph after the given paragraph."""
    new_p = docx.oxml.OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    # Create a temporary paragraph object to add run with image
    from docx.text.paragraph import Paragraph
    temp_p = Paragraph(new_p, paragraph._element.getparent())
    run = temp_p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    return temp_p

def replace_paragraph_with_image(paragraph, img_path, width_inches=5.5):
    """Replace paragraph content with an image."""
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].add_picture(img_path, width=Inches(width_inches))
    else:
        run = paragraph.add_run()
        run.add_picture(img_path, width=Inches(width_inches))

def replace_paragraph_with_code(paragraph, code_text):
    """Replace paragraph text with code content."""
    for run in paragraph.runs:
        run.text = ""
    lines = code_text.strip().split('\n')
    if paragraph.runs:
        paragraph.runs[0].text = lines[0]
        paragraph.runs[0].font.name = "Courier New"
        paragraph.runs[0].font.size = Pt(8)
    else:
        run = paragraph.add_run(lines[0])
        run.font.name = "Courier New"
        run.font.size = Pt(8)
    
    insert_point = paragraph._element
    for line in lines[1:]:
        new_p = docx.oxml.OxmlElement('w:p')
        new_r = docx.oxml.OxmlElement('w:r')
        new_rpr = docx.oxml.OxmlElement('w:rPr')
        new_rfont = docx.oxml.OxmlElement('w:rFonts')
        new_rfont.set(qn('w:ascii'), 'Courier New')
        new_rfont.set(qn('w:hAnsi'), 'Courier New')
        new_rpr.append(new_rfont)
        new_sz = docx.oxml.OxmlElement('w:sz')
        new_sz.set(qn('w:val'), '16')
        new_rpr.append(new_sz)
        new_r.append(new_rpr)
        new_t = docx.oxml.OxmlElement('w:t')
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = line
        new_r.append(new_t)
        new_p.append(new_r)
        insert_point.addnext(new_p)
        insert_point = new_p

# ============================================================
# IMAGE MAPPING - which image goes where
# ============================================================
# Based on document analysis, map paragraph text patterns to images
IMAGE_MAP = {
    # Architecture diagram (empty paragraphs after "Arsitektur Sistem")
    'Arsitektur Sistem': {'img': img_arch, 'width': 5.5},
    # Use Case
    'Use Case Diagram': {'img': img_uc, 'width': 5.5},
    # Editorial workflow
    'Alur Kerja Editorial': {'img': img_ew, 'width': 5.5},
    # Class diagrams
    'Class Diagram-Backend': {'img': img_cb, 'width': 5.5},
    'Conceptual Data Model': {'img': img_cdm, 'width': 5.5},
}

ACTIVITY_MAP = {
    'Proses Login & Autentikasi': act_imgs['login'],
    'Silent Token Refresh': act_imgs['refresh'],
    'Alur Penulisan & Publikasi Artikel': act_imgs['article'],
    'Moderasi Komentar': act_imgs['comment'],
    'Scheduled Publishing': act_imgs['scheduled'],
    'Upload Media': act_imgs['upload'],
    'Bookmark Artikel': act_imgs['bookmark'],
    'Pencarian Artikel': act_imgs['search'],
    'Logout & Invalidasi Sesi': act_imgs['logout'],
}

SEQUENCE_MAP = {
    'Sequence Diagram-Proses Login': seq_imgs['seq_login'],
    'Sequence Diagram-Silent Token Refresh': seq_imgs['seq_refresh'],
    'Sequence Diagram-CRUD Artikel CMS': seq_imgs['seq_crud'],
    'Sequence Diagram-Review dan Publish Artikel': seq_imgs['seq_review'],
    'Sequence Diagram-Bookmart Artikel': seq_imgs['seq_bookmark'],
    'Sequence Diagram-Page View Analytics': seq_imgs['seq_analytics'],
    'Sequence Diagram-Registrasi Akun Baru': seq_imgs['seq_register'],
    'Sequence Diagram-Ambil Daftar Artikel': seq_imgs['seq_list'],
    'Sequence Diagram-Submit Artikel Baru': seq_imgs['seq_submit'],
    'Sequence Diagram-Post Komentar': seq_imgs['seq_comment'],
    'Sequence Diagram-Upload Media': seq_imgs['seq_upload'],
    'Sequence Diagram-Edit dan Update Artikel': seq_imgs['seq_edit'],
}

# Process: find empty paragraphs after diagram labels and insert images
filled_imgs = 0
i = 0
while i < len(paragraphs):
    text = paragraphs[i].text.strip()
    
    # Check if this is a diagram label followed by empty paragraph
    for label, info in IMAGE_MAP.items():
        if label in text and i+1 < len(paragraphs) and paragraphs[i+1].text.strip() == '':
            replace_paragraph_with_image(paragraphs[i+1], info['img'], info.get('width', 5.5))
            filled_imgs += 1
            print(f"  Inserted image for: {label}")
            break
    
    # Activity diagrams
    for label, img in ACTIVITY_MAP.items():
        if label in text:
            # Find next empty paragraph
            for j in range(i+1, min(i+5, len(paragraphs))):
                if paragraphs[j].text.strip() == '':
                    replace_paragraph_with_image(paragraphs[j], img, 5.5)
                    filled_imgs += 1
                    print(f"  Inserted activity image: {label}")
                    break
            break
    
    # Sequence diagrams
    for label, img in SEQUENCE_MAP.items():
        if label in text:
            for j in range(i+1, min(i+5, len(paragraphs))):
                if paragraphs[j].text.strip() == '':
                    replace_paragraph_with_image(paragraphs[j], img, 5.0)
                    filled_imgs += 1
                    print(f"  Inserted sequence image: {label}")
                    break
            break
    
    # GAMBAR placeholders
    if text == 'GAMBAR':
        # Determine which diagram based on context
        context = ''
        for k in range(max(0, i-10), i):
            context += paragraphs[k].text + ' '
        
        img_to_use = None
        if 'Arsitektur' in context or 'Design Pattern' in context:
            img_to_use = img_sr
        elif 'Buka Daftar' in context or 'Ambil Daftar' in context:
            img_to_use = seq_imgs['seq_list']
        elif 'Submit Artikel' in context:
            img_to_use = seq_imgs['seq_submit']
        elif 'Post Komentar' in context:
            img_to_use = seq_imgs['seq_comment']
        elif 'Upload Media' in context:
            img_to_use = seq_imgs['seq_upload']
        elif 'Edit dan Update' in context:
            img_to_use = seq_imgs['seq_edit']
        elif 'Struktur Folder Backend' in context:
            img_to_use = img_bf
        
        if img_to_use:
            replace_paragraph_with_image(paragraphs[i], img_to_use, 5.0)
            filled_imgs += 1
            print(f"  Replaced GAMBAR placeholder at paragraph {i}")
    
    # Empty paragraph after "Struktur Folder Proyek"
    if 'Struktur Folder Proyek' in text and i+1 < len(paragraphs):
        if paragraphs[i+1].text.strip() == '':
            replace_paragraph_with_image(paragraphs[i+1], img_ff, 4.5)
            filled_imgs += 1
            print(f"  Inserted frontend folder image")
    
    # Empty paragraph after backend folder label
    if 'Struktur Folder Backend' in text and i+1 < len(paragraphs):
        if paragraphs[i+1].text.strip() == '':
            replace_paragraph_with_image(paragraphs[i+1], img_bf, 4.5)
            filled_imgs += 1
            print(f"  Inserted backend folder image")
    
    i += 1

# ============================================================
# FILL KODE PLACEHOLDERS
# ============================================================

# Code snippets mapped to actual source files
CODE_SNIPPETS = {
    'getBookmarks': '''use Illuminate\\Support\\Facades\\DB;
use App\\Models\\Article;

class BookmarkService
{
    public function getBookmarks(int $userId, int $page = 1, int $limit = 10): array
    {
        $offset = ($page - 1) * $limit;
        $total = DB::table('bookmarks')->where('user_id', $userId)->count();

        $articleIds = DB::table('bookmarks')
            ->where('user_id', $userId)
            ->orderBy('created_at', 'desc')
            ->skip($offset)->take($limit)
            ->pluck('article_id')->toArray();

        $articles = Article::with(['category', 'author', 'tags'])
            ->whereIn('id', $articleIds)
            ->where('status', 'published')
            ->get()->map(fn($a) => [
                'id' => $a->id, 'title' => $a->title,
                'slug' => $a->slug, 'cover_url' => $a->cover_url,
                'category' => $a->category?->name,
                'author' => $a->author?->name,
            ]);

        return ['data' => $articles, 'total' => $total,
                'page' => $page, 'total_pages' => ceil($total / $limit)];
    }
}''',
    'bookmark_controller': '''<?php
namespace App\\Http\\Controllers;
use App\\Services\\BookmarkService;
use Illuminate\\Http\\Request;

class BookmarkController extends Controller
{
    public function __construct(private BookmarkService $bookmarkService) {}

    public function index(Request $request)
    {
        $userId = auth('api')->id();
        $bookmarks = $this->bookmarkService->getBookmarks(
            $userId, $request->input('page', 1), $request->input('limit', 10)
        );
        return response()->json(['status' => 'success', 'data' => $bookmarks]);
    }

    public function toggle(Request $request)
    {
        $request->validate(['article_id' => 'required|string|exists:articles,id']);
        $userId = auth('api')->id();
        $isBookmarked = $this->bookmarkService->toggle($userId, $request->article_id);
        return response()->json(['status' => 'success', 'data' => ['bookmarked' => $isBookmarked]]);
    }
}''',
    'dashboard_index': '''<?php
namespace App\\Http\\Controllers;
use Illuminate\\Support\\Facades\\DB;

class CmsDashboardController extends Controller
{
    public function index()
    {
        $user = auth('api')->user();
        $stats = $user->role === 'admin'
            ? $this->getAdminStats()
            : $this->getEditorStats($user->id, $user->role);
        return response()->json(['status' => 'success', 'data' => $stats]);
    }

    private function getAdminStats(): array
    {
        return [
            'type' => 'admin',
            'total_views' => DB::table('page_views')->count(),
            'total_berita' => DB::table('articles')->count(),
            'new_users' => DB::table('users')
                ->whereMonth('created_at', date('m'))->count(),
        ];
    }
}''',
    'editor_stats': '''private function getEditorStats(string $userId, string $role): array
{
    $query = DB::table('articles');
    if ($role === 'journalist') {
        $query->where('author_id', $userId);
    } else {
        $query->where(fn($q) => $q->where('author_id', $userId)
                                  ->orWhere('published_by', $userId));
    }

    return [
        'type' => $role,
        'published' => (clone $query)->where('status', 'published')->count(),
        'draft' => (clone $query)->where('status', 'draft')->count(),
        'review' => (clone $query)->where('status', 'review')->count(),
        'yearly_data' => DB::table('articles')
            ->select(DB::raw('MONTH(created_at) as month'), DB::raw('COUNT(*) as total'))
            ->whereYear('created_at', date('Y'))->groupBy('month')->get(),
        'category_data' => DB::table('articles')->join('categories', 'articles.category_id', '=', 'categories.id')
            ->select('categories.name', DB::raw('COUNT(*) as total'))->groupBy('categories.id')->get(),
    ];
}''',
    'api_routes': '''<?php
use Illuminate\\Support\\Facades\\Route;

Route::prefix('auth')->group(function () {
    Route::post('/register', [AuthController::class, 'register']);
    Route::post('/login', [AuthController::class, 'login']);
    Route::post('/refresh', [AuthController::class, 'refresh']);
});

Route::middleware('auth:api')->group(function () {
    Route::post('/auth/logout', [AuthController::class, 'logout']);
    Route::get('/auth/me', [AuthController::class, 'me']);
    Route::get('/articles', [ArticleController::class, 'index']);
    Route::get('/articles/{slug}', [ArticleController::class, 'showBySlug']);
    Route::get('/articles/{id}/comments', [CommentController::class, 'index']);
    Route::post('/articles/{id}/comments', [CommentController::class, 'store']);
    Route::get('/bookmarks', [BookmarkController::class, 'index']);
    Route::post('/bookmarks', [BookmarkController::class, 'toggle']);
    Route::post('/media', [MediaController::class, 'store']);
    Route::prefix('cms')->group(function () {
        Route::get('/statistics', [CmsDashboardController::class, 'index']);
        Route::get('/articles', [CmsArticleController::class, 'index']);
        Route::post('/articles', [CmsArticleController::class, 'store']);
        Route::patch('/articles/{id}/status', [CmsArticleController::class, 'updateStatus']);
        Route::put('/articles/{id}', [CmsArticleController::class, 'update']);
        Route::delete('/articles/{id}', [CmsArticleController::class, 'destroy']);
    });
    Route::middleware('admin')->group(function () {
        Route::apiResource('/users', UserController::class);
    });
});''',
    'middleware_admin': '''<?php
namespace App\\Http\\Middleware;
use Closure;
use Illuminate\\Http\\Request;

class EnsureAdmin
{
    public function handle(Request $request, Closure $next)
    {
        $user = auth('api')->user();
        if (!$user || $user->role !== 'admin') {
            return response()->json([
                'status' => 'error', 'error' => 'FORBIDDEN',
                'message' => 'Akses ditolak. Hanya admin.',
            ], 403);
        }
        return $next($request);
    }
}''',
    'migration_articles': '''<?php
use Illuminate\\Database\\Migrations\\Migration;
use Illuminate\\Database\\Schema\\Blueprint;
use Illuminate\\Support\\Facades\\Schema;

return new class extends Migration
{
    public function up(): void
    {
        Schema::create('articles', function (Blueprint $table) {
            $table->uuid('id')->primary();
            $table->string('title');
            $table->string('slug')->unique();
            $table->text('excerpt')->nullable();
            $table->longText('content');
            $table->string('cover_url')->nullable();
            $table->enum('status', ['draft','review','published','scheduled','rejected','archived'])->default('draft');
            $table->unsignedBigInteger('view_count')->default(0);
            $table->boolean('is_featured')->default(false);
            $table->uuid('author_id');
            $table->uuid('category_id')->nullable();
            $table->uuid('published_by')->nullable();
            $table->uuid('locked_by')->nullable();
            $table->timestamp('published_at')->nullable();
            $table->timestamps();
            $table->foreign('author_id')->references('id')->on('users')->cascadeOnDelete();
            $table->foreign('category_id')->references('id')->on('categories')->nullOnDelete();
            $table->index('status');
            $table->index('view_count');
        });
    }
};''',
    'user_seeder': '''<?php
namespace Database\\Seeders;
use Illuminate\\Database\\Seeder;
use Illuminate\\Support\\Facades\\DB;
use Illuminate\\Support\\Facades\\Hash;

class UserSeeder extends Seeder
{
    public function run(): void
    {
        $users = json_decode(file_get_contents(database_path('data/dummy2.json')), true);
        foreach ($users as $user) {
            DB::table('users')->insert([
                'id' => $user['id'], 'name' => $user['name'],
                'email' => $user['email'],
                'password' => Hash::make($user['password']),
                'role' => $user['role'],
                'is_active' => $user['is_active'] ?? true,
                'created_at' => now(), 'updated_at' => now(),
            ]);
        }
    }
}''',
    'scheduled_publish': '''<?php
namespace App\\Console\\Commands;
use App\\Services\\ScheduledPublishService;
use Illuminate\\Console\\Command;

class PublishScheduledArticles extends Command
{
    protected $signature = 'articles:publish-scheduled';
    protected $description = 'Publish articles that reached their scheduled time';

    public function __construct(private ScheduledPublishService $service)
    {
        parent::__construct();
    }

    public function handle(): int
    {
        $count = $this->service->publishDueArticles();
        $this->info("Published {$count} scheduled article(s).");
        return Command::SUCCESS;
    }
}''',
    'frontend_folder': '''frontend/
\\u251c\\u2500\\u2500 app/
\\u2502   \\u251c\\u2500\\u2500 (main)/           # Portal Publik
\\u2502   \\u2502   \\u251c\\u2500\\u2500 page.tsx      # Dashboard Reader
\\u2502   \\u2502   \\u251c\\u2500\\u2500 [slug]/       # Detail Artikel (SSG)
\\u2502   \\u2502   \\u2514\\u2500\\u2500 bookmark/
\\u2502   \\u251c\\u2500\\u2500 cms/              # CMS Redaksi
\\u2502   \\u2502   \\u251c\\u2500\\u2500 dashboard/
\\u2502   \\u2502   \\u251c\\u2500\\u2500 artikel/tulis/
\\u2502   \\u2502   \\u2514\\u2500\\u2500 artikel/tinjauan/
\\u2502   \\u251c\\u2500\\u2500 components/
\\u2502   \\u251c\\u2500\\u2500 constants/
\\u2502   \\u2514\\u2500\\u2500 types/index.ts
\\u2502   \\u2514\\u2500\\u2500 layout.tsx
\\u251c\\u2500\\u2500 lib/
\\u2502   \\u251c\\u2500\\u2500 api/ (articles, bookmarks)
\\u2502   \\u251c\\u2500\\u2500 axios.ts (interceptors)
\\u2502   \\u2514\\u2500\\u2500 validations.ts (Zod)
\\u251c\\u2500\\u2500 stores/authStore.ts (Zustand)
\\u2514\\u2500\\u2500 utils/cropImage.ts''',
    'dashboard_page': '''"use client";
import { useEffect, useState } from "react";
import { useAuthStore } from "@/stores/authStore";
import axiosInstance from "@/lib/axios";
import { BarChart, Bar, XAxis, YAxis, PieChart, Pie, Cell } from "recharts";

export default function EditorDashboardView() {
  const { user } = useAuthStore();
  const [stats, setStats] = useState<any>(null);

  useEffect(() => {
    axiosInstance.get("/api/cms/statistics").then((res) => setStats(res.data.data));
  }, []);

  if (!stats) return <div>Loading...</div>;

  return (
    <div className="p-6 space-y-6">
      <h1 className="text-2xl font-bold">Dashboard — {user?.name}</h1>
      <div className="grid grid-cols-3 gap-4">
        <StatCard title="Published" value={stats.published} />
        <StatCard title="Draft" value={stats.draft} />
        <StatCard title="Review" value={stats.review} />
      </div>
      <BarChart data={stats.yearly_data}>
        <Bar dataKey="total" fill="#3B82F6" />
      </BarChart>
    </div>
  );
}''',
    'bookmark_api': '''import axiosInstance from "../axios";
import type { ApiSuccess, Bookmark, Pagination } from "@/app/types";

export async function getBookmarks(page = 1, limit = 10) {
  const res = await axiosInstance.get<ApiSuccess<Bookmark[] & Pagination>>(
    `/api/bookmarks?page=${page}&limit=${limit}`
  );
  return res.data.data;
}

export async function toggleBookmark(articleId: string) {
  const res = await axiosInstance.post<ApiSuccess<{ bookmarked: boolean }>>(
    "/api/bookmarks", { article_id: articleId }
  );
  return res.data.data;
}''',
    'axios_interceptor': '''import axios from "axios";
import { getAccessToken, setAccessToken, useAuthStore } from "@/stores/authStore";

const axiosInstance = axios.create({
  baseURL: process.env.NEXT_PUBLIC_API_URL,
});

axiosInstance.interceptors.request.use((config) => {
  const token = getAccessToken();
  if (token) config.headers.Authorization = `Bearer ${token}`;
  return config;
});

axiosInstance.interceptors.response.use(
  (response) => response,
  async (error) => {
    const originalRequest = error.config;
    if (error.response?.data?.error === "TOKEN_EXPIRED" && !originalRequest._retry) {
      originalRequest._retry = true;
      try {
        const refreshToken = localStorage.getItem("refresh_token");
        const res = await axios.post(`${process.env.NEXT_PUBLIC_API_URL}/api/auth/refresh`, { refresh_token: refreshToken });
        const newToken = res.data.data.access_token;
        setAccessToken(newToken);
        originalRequest.headers.Authorization = `Bearer ${newToken}`;
        return axiosInstance(originalRequest);
      } catch {
        useAuthStore.getState().logout();
        window.location.href = "/login";
      }
    }
    return Promise.reject(error);
  }
);
export default axiosInstance;''',
    'typescript_interfaces': '''export interface User {
  id: string; name: string; email: string;
  role: "reader" | "journalist" | "editor" | "admin";
  is_active: boolean; created_at: string;
}
export interface Article {
  id: string; title: string; slug: string;
  excerpt: string | null; content: string;
  cover_url: string | null;
  status: "draft" | "review" | "published" | "scheduled" | "rejected" | "archived";
  view_count: number; is_featured: boolean;
  author: Pick<User, "id" | "name">;
  category: { id: string; name: string; slug: string } | null;
  tags: { id: string; name: string; slug: string }[];
  published_at: string | null; created_at: string;
}
export interface Comment {
  id: string; content: string;
  status: "pending" | "approved" | "rejected";
  user: Pick<User, "id" | "name">;
  parent_id: string | null; replies?: Comment[];
}
export interface Pagination { total: number; page: number; limit: number; total_pages: number; }
export interface ApiSuccess<T> { status: "success"; data: T; }
export interface ApiError { status: "error"; error: string; message: string; }''',
    'zod_schemas': '''import { z } from "zod";

export const loginSchema = z.object({
  email: z.string().email("Email tidak valid"),
  password: z.string().min(6, "Password minimal 6 karakter"),
});

export const registerSchema = z.object({
  name: z.string().min(3).max(100),
  email: z.string().email(),
  password: z.string().min(6),
  password_confirmation: z.string(),
}).refine((d) => d.password === d.password_confirmation, {
  message: "Konfirmasi password tidak cocok",
  path: ["password_confirmation"],
});

export const articleSchema = z.object({
  title: z.string().min(5).max(255),
  content: z.string().min(10),
  category_id: z.string().uuid(),
  tag_ids: z.array(z.string().uuid()).default([]),
});

export const scheduleSchema = z.object({
  scheduled_at: z.string().refine((val) => {
    const diff = (new Date(val).getTime() - Date.now()) / 60000;
    return diff >= 5;
  }, "Waktu tayang minimal 5 menit dari sekarang"),
});

export const commentSchema = z.object({
  content: z.string().min(1).max(1000),
  parent_id: z.string().uuid().nullable().optional(),
});''',
    'roles_constants': '''// constants/roles.ts
export const ROLES = {
  READER: "reader", JOURNALIST: "journalist",
  EDITOR: "editor", ADMIN: "admin",
} as const;

export function canPublish(role: string): boolean {
  return role === ROLES.EDITOR || role === ROLES.ADMIN;
}
export function canDeleteComment(role: string): boolean {
  return role === ROLES.EDITOR || role === ROLES.ADMIN;
}
export function canManageUsers(role: string): boolean {
  return role === ROLES.ADMIN;
}''',
    'articles_api': '''import axiosInstance from "../axios";
import type { ApiSuccess, Article, Comment, Pagination } from "@/app/types";

export async function getArticles(page = 1, limit = 10) {
  const res = await axiosInstance.get<ApiSuccess<Article[] & Pagination>>(`/api/articles?page=${page}&limit=${limit}`);
  return res.data.data;
}
export async function getCmsArticles(page = 1, limit = 10, filters?: Record<string, string>) {
  const params = new URLSearchParams({ page: String(page), limit: String(limit) });
  if (filters) Object.entries(filters).forEach(([k, v]) => params.append(k, v));
  const res = await axiosInstance.get<ApiSuccess<Article[] & Pagination>>(`/api/cms/articles?${params}`);
  return res.data.data;
}
export async function createArticle(data: FormData) {
  const res = await axiosInstance.post<ApiSuccess<Article>>("/api/cms/articles", data);
  return res.data.data;
}
export async function updateArticleStatus(id: string, status: string, note?: string) {
  const res = await axiosInstance.patch<ApiSuccess<Article>>(`/api/cms/articles/${id}/status`, { status, change_note: note });
  return res.data.data;
}
export async function lockArticle(id: string) { await axiosInstance.post(`/api/cms/articles/${id}/lock`); }
export async function unlockArticle(id: string) { await axiosInstance.post(`/api/cms/articles/${id}/unlock`); }''',
    'auth_store': '''import { create } from "zustand";
import { persist } from "zustand/middleware";
import type { User } from "@/app/types";

interface AuthState {
  accessToken: string | null; user: User | null; isAuthenticated: boolean;
  setAuth: (token: string, user: User) => void;
  setAccessToken: (token: string) => void;
  logout: () => void;
}

export const useAuthStore = create<AuthState>()(persist((set) => ({
  accessToken: null, user: null, isAuthenticated: false,
  setAuth: (token, user) => set({ accessToken: token, user, isAuthenticated: true }),
  setAccessToken: (token) => set({ accessToken: token }),
  logout: () => {
    localStorage.removeItem("refresh_token");
    set({ accessToken: null, user: null, isAuthenticated: false });
  },
}), { name: "auth-storage" }));

export const getAccessToken = () => useAuthStore.getState().accessToken;
export const setAccessToken = (token: string) => useAuthStore.getState().setAccessToken(token);''',
}

# Order of KODE placeholders in document
KODE_ORDER = [
    'getBookmarks', 'bookmark_controller', 'dashboard_index', 'editor_stats',
    'api_routes', 'middleware_admin', 'migration_articles', 'user_seeder',
    'scheduled_publish', 'frontend_folder', 'dashboard_page', 'dashboard_page',
    'bookmark_api', 'axios_interceptor', 'typescript_interfaces', 'zod_schemas',
    'roles_constants', 'articles_api', 'auth_store',
]

# Fill KODE placeholders
kode_filled = 0
for i, p in enumerate(paragraphs):
    if p.text.strip() == 'KODE' and kode_filled < len(KODE_ORDER):
        code_key = KODE_ORDER[kode_filled]
        code_text = CODE_SNIPPETS.get(code_key, '// Code not found')
        replace_paragraph_with_code(p, code_text)
        kode_filled += 1
        print(f"  Filled KODE #{kode_filled}: {code_key}")

# Save
output = "23081010208_laporan_pkl_filled.docx"
doc.save(output)
print(f"\nDone! Filled {filled_imgs} diagram images + {kode_filled} code snippets")
print(f"Saved to: {output}")
