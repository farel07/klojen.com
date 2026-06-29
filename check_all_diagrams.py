from docx import Document
doc = Document('laporan_pkl.docx')
NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# Show ALL content around the problematic diagram areas
# Check Moderasi Komentar (around 471-476)
print('=== Activity Diagram — Moderasi Komentar area ===')
for i in range(469, 480):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        drawings = p._element.findall(f'.//{NS}drawing')
        has_img = len(drawings) > 0
        label = 'IMG' if has_img else 'TXT'
        text = t[:200] if t else '(empty)'
        print(f'[{i}] [{label}] {text}')

print('\n=== Sequence Diagram — CRUD Artikel CMS area ===')
for i in range(501, 510):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        drawings = p._element.findall(f'.//{NS}drawing')
        has_img = len(drawings) > 0
        label = 'IMG' if has_img else 'TXT'
        text = t[:200] if t else '(empty)'
        print(f'[{i}] [{label}] {text}')

# Also check ALL diagram title paragraphs and their following paragraphs
print('\n=== ALL Diagram titles and their next paragraphs ===')
for i in range(455, 545):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        drawings = p._element.findall(f'.//{NS}drawing')
        has_img = len(drawings) > 0
        if t and ('Diagram' in t or 'diagram' in t) and ('Activity' in t or 'Sequence' in t or 'Class' in t or 'Use Case' in t or 'State' in t):
            # This is a diagram title - check next paragraph
            next_i = i + 1
            if next_i < len(doc.paragraphs):
                next_p = doc.paragraphs[next_i]
                next_drawings = next_p._element.findall(f'.//{NS}drawing')
                next_has_img = len(next_drawings) > 0
                next_t = next_p.text.strip()[:80] if next_p.text.strip() else '(empty)'
                status = 'HAS_IMG' if next_has_img else 'NO_IMG'
                print(f'[{i}] {t[:80]}')
                print(f'  [{next_i}] [{status}] {next_t}')
