from docx import Document
from lxml import etree
doc = Document('laporan_pkl.docx')

# Print all namespaces used in drawings
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
    if not drawings:
        # Try without namespace
        drawings = p._element.findall('.//drawing')
    
    # Check for any image reference
    xml_str = etree.tostring(p._element, encoding='unicode')
    if 'blipFill' in xml_str or 'blip' in xml_str:
        # Extract embed ID
        import re
        embeds = re.findall(r'r:embed="([^"]+)"', xml_str)
        title = ''
        for j in range(i-1, max(0, i-5), -1):
            t = doc.paragraphs[j].text.strip()
            if t:
                title = t[:80]
                break
        if embeds:
            print(f'[{i}] embeds={embeds}, title: {title}')
