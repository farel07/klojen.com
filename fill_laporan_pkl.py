"""Fill all empty KODE and GAMBAR placeholders in 23081010208_laporan pkl.docx"""
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

doc = docx.Document('23081010208_laporan pkl.docx')

# ============================================================
# CODE SNIPPETS - mapped to actual source files
# ============================================================

CODE_BOOKMARK_SERVICE_TOGGLE = '''use Illuminate\\Support\\Facades\\DB;
use Illuminate\\Support\\Str;

class BookmarkService
{
    /**
     * Toggle bookmark: add if not exists, remove if exists
     */
    public function toggle(int $userId, string $articleId): bool
    {
        $existing = DB::table('bookmarks')
            ->where('user_id', $userId)
            ->where('article_id', $articleId)
            ->first();

        if ($existing) {
            DB::table('bookmarks')
                ->where('user_id', $userId)
                ->where('article_id', $articleId)
                ->delete();
            return false;
        }

        DB::table('bookmarks')->insert([
            'id'         => Str::uuid()->toString(),
            'user_id'    => $userId,
            'article_id' => $articleId,
            'created_at' => now(),
        ]);
        return true;
    }

    /**
     * Get paginated bookmarks for a user with article details
     */
    public function getBookmarks(int $userId, int $page = 1, int $limit = 10): array
    {
        $offset = ($page - 1) * $limit;

        $total = DB::table('bookmarks')
            ->where('user_id', $userId)
            ->count();

        $bookmarkArticleIds = DB::table('bookmarks')
            ->where('user_id', $userId)
            ->orderBy('created_at', 'desc')
            ->skip($offset)
            ->take($limit)
            ->pluck('article_id')
            ->toArray();

        $articles = [];
        if (!empty($bookmarkArticleIds)) {
            $articles = \\App\\Models\\Article::with(['category', 'author', 'tags'])
                ->whereIn('id', $bookmarkArticleIds)
                ->where('status', 'published')
                ->get()
                ->keyBy('id')
                ->map(function ($article) {
                    return [
                        'id'         => $article->id,
                        'title'      => $article->title,
                        'slug'       => $article->slug,
                        'excerpt'    => $article->excerpt,
                        'cover_url'  => $article->cover_url,
                        'category'   => $article->category ? $article->category->name : null,
                        'author'     => $article->author ? $article->author->name : null,
                        'created_at' => $article->created_at->toIso8601String(),
                    ];
                })->toArray();
        }

        return [
            'data'        => array_values($articles),
            'total'       => $total,
            'page'        => $page,
            'limit'       => $limit,
            'total_pages' => ceil($total / $limit),
        ];
    }
}'''

CODE_BOOKMARK_CONTROLLER = '''<?php

namespace App\\Http\\Controllers;

use App\\Services\\BookmarkService;
use Illuminate\\Http\\Request;

class BookmarkController extends Controller
{
    private BookmarkService $bookmarkService;

    public function __construct(BookmarkService $bookmarkService)
    {
        $this->bookmarkService = $bookmarkService;
    }

    public function index(Request $request)
    {
        $userId = auth('api')->id();
        $page   = $request->input('page', 1);
        $limit  = $request->input('limit', 10);

        $bookmarks = $this->bookmarkService->getBookmarks($userId, $page, $limit);

        return response()->json([
            'status' => 'success',
            'data'   => $bookmarks,
        ]);
    }

    public function toggle(Request $request)
    {
        $request->validate([
            'article_id' => 'required|string|exists:articles,id',
        ]);

        $userId = auth('api')->id();
        $articleId = $request->input('article_id');

        $isBookmarked = $this->bookmarkService->toggle($userId, $articleId);

        return response()->json([
            'status' => 'success',
            'data'   => ['bookmarked' => $isBookmarked],
        ]);
    }
}'''

CODE_CMS_DASHBOARD_INDEX = '''<?php

namespace App\\Http\\Controllers;

use Illuminate\\Support\\Facades\\DB;

class CmsDashboardController extends Controller
{
    public function index()
    {
        $user = auth('api')->user();
        $role = $user->role;

        if ($role === 'admin') {
            $stats = $this->getAdminStats();
        } else {
            $stats = $this->getEditorStats($user->id, $role);
        }

        return response()->json([
            'status' => 'success',
            'data'   => $stats,
        ]);
    }

    private function getAdminStats(): array
    {
        $totalPageViews = DB::table('page_views')->count();
        $totalBerita    = DB::table('articles')->count();
        $newUsers       = DB::table('users')
            ->whereMonth('created_at', date('m'))
            ->whereYear('created_at', date('Y'))
            ->count();

        // Sparkline: page views per hari (7 hari terakhir)
        $sparklines = DB::table('page_views')
            ->select(DB::raw('DATE(created_at) as date'), DB::raw('COUNT(*) as views'))
            ->where('created_at', '>=', now()->subDays(7))
            ->groupBy('date')
            ->orderBy('date')
            ->get();

        // Visitor data: top 5 artikel paling banyak dilihat
        $topArticles = DB::table('articles')
            ->select('title', 'view_count')
            ->orderByDesc('view_count')
            ->take(5)
            ->get();

        return [
            'type'         => 'admin',
            'total_views'  => $totalPageViews,
            'total_berita' => $totalBerita,
            'new_users'    => $newUsers,
            'sparklines'   => $sparklines,
            'top_articles' => $topArticles,
        ];
    }'''

CODE_EDITOR_STATS = '''    private function getEditorStats(string $userId, string $role): array
    {
        $query = DB::table('articles');

        if ($role === 'journalist') {
            $query->where('author_id', $userId);
        } else {
            // editor: see own articles + articles they published
            $query->where(function ($q) use ($userId) {
                $q->where('author_id', $userId)
                  ->orWhere('published_by', $userId);
            });
        }

        $publishedCount = (clone $query)->where('status', 'published')->count();
        $draftCount     = (clone $query)->where('status', 'draft')->count();
        $reviewCount    = (clone $query)->where('status', 'review')->count();

        // Yearly data for BarChart
        $yearlyData = DB::table('articles')
            ->select(DB::raw('MONTH(created_at) as month'), DB::raw('COUNT(*) as total'))
            ->whereYear('created_at', date('Y'))
            ->groupBy('month')
            ->orderBy('month')
            ->get();

        // Category distribution for PieChart
        $categoryData = DB::table('articles')
            ->join('categories', 'articles.category_id', '=', 'categories.id')
            ->select('categories.name', DB::raw('COUNT(*) as total'))
            ->groupBy('categories.id', 'categories.name')
            ->get();

        return [
            'type'           => $role,
            'published'      => $publishedCount,
            'draft'          => $draftCount,
            'review'         => $reviewCount,
            'yearly_data'    => $yearlyData,
            'category_data'  => $categoryData,
        ];
    }
}'''

CODE_API_ROUTES = '''<?php

use Illuminate\\Support\\Facades\\Route;

// === Auth ===
Route::prefix('auth')->group(function () {
    Route::post('/register', [AuthController::class, 'register']);
    Route::post('/login', [AuthController::class, 'login']);
    Route::post('/refresh', [AuthController::class, 'refresh']);
    Route::post('/forgot-password', [AuthController::class, 'forgotPassword']);
    Route::post('/reset-password', [AuthController::class, 'resetPassword']);
});

// === Protected Routes (JWT) ===
Route::middleware('auth:api')->group(function () {

    // Auth
    Route::post('/auth/logout', [AuthController::class, 'logout']);
    Route::get('/auth/me', [AuthController::class, 'me']);
    Route::put('/auth/profile', [AuthController::class, 'updateProfile']);
    Route::put('/auth/change-password', [AuthController::class, 'changePassword']);

    // Articles (public portal)
    Route::get('/articles', [ArticleController::class, 'index']);
    Route::get('/articles/{slug}', [ArticleController::class, 'showBySlug']);
    Route::get('/articles/{id}/related', [ArticleController::class, 'related']);

    // Comments
    Route::get('/articles/{articleId}/comments', [CommentController::class, 'index']);
    Route::post('/articles/{articleId}/comments', [CommentController::class, 'store']);
    Route::delete('/comments/{id}', [CommentController::class, 'destroy']);

    // Bookmarks
    Route::get('/bookmarks', [BookmarkController::class, 'index']);
    Route::post('/bookmarks', [BookmarkController::class, 'toggle']);

    // Media
    Route::get('/media', [MediaController::class, 'index']);
    Route::post('/media', [MediaController::class, 'store']);
    Route::delete('/media/{id}', [MediaController::class, 'destroy']);

    // Analytics
    Route::post('/analytics/track', [AnalyticsController::class, 'track']);

    // CMS Articles
    Route::prefix('cms')->group(function () {
        Route::get('/statistics', [CmsDashboardController::class, 'index']);
        Route::get('/articles', [CmsArticleController::class, 'index']);
        Route::post('/articles', [CmsArticleController::class, 'store']);
        Route::get('/articles/{id}', [CmsArticleController::class, 'show']);
        Route::patch('/articles/{id}/status', [CmsArticleController::class, 'updateStatus']);
        Route::put('/articles/{id}', [CmsArticleController::class, 'update']);
        Route::delete('/articles/{id}', [CmsArticleController::class, 'destroy']);
        Route::post('/articles/{id}/lock', [CmsArticleController::class, 'lock']);
        Route::post('/articles/{id}/unlock', [CmsArticleController::class, 'unlock']);

        // CMS Comments
        Route::get('/comments', [CmsCommentController::class, 'index']);
        Route::patch('/comments/{id}/status', [CmsCommentController::class, 'updateStatus']);

        // CMS Categories & Tags
        Route::apiResource('/categories', CmsCategoryController::class);
        Route::apiResource('/tags', CmsTagController::class);
    });

    // Users (admin only)
    Route::middleware('admin')->group(function () {
        Route::get('/users', [UserController::class, 'index']);
        Route::post('/users', [UserController::class, 'store']);
        Route::get('/users/{id}', [UserController::class, 'show']);
        Route::put('/users/{id}', [UserController::class, 'update']);
        Route::delete('/users/{id}', [UserController::class, 'destroy']);
    });
});'''

CODE_MIDDLEWARE_ADMIN = '''<?php

namespace App\\Http\\Middleware;

use Closure;
use Illuminate\\Http\\Request;

class EnsureAdmin
{
    public function handle(Request $request, Closure $next)
    {
        $user = auth('api')->user();

        if (!$user || $user->role !== 'admin') {
            return response()->json([
                'status'  => 'error',
                'error'   => 'FORBIDDEN',
                'message' => 'Akses ditolak. Hanya admin yang dapat mengakses resource ini.',
            ], 403);
        }

        return $next($request);
    }
}'''

CODE_ARTICLES_MIGRATION = '''<?php

use Illuminate\\Database\\Migrations\\Migration;
use Illuminate\\Database\\Schema\\Blueprint;
use Illuminate\\Support\\Facades\\Schema;

return new class extends Migration
{
    public function up(): void
    {
        Schema::create('articles', function (Blueprint $table) {
            $table->uuid('id')->primary();
            $table->string('title');
            $table->string('slug')->unique();
            $table->text('excerpt')->nullable();
            $table->longText('content');
            $table->string('cover_url')->nullable();
            $table->enum('status', ['draft', 'review', 'published', 'scheduled', 'rejected', 'archived'])
                  ->default('draft');
            $table->unsignedBigInteger('view_count')->default(0);
            $table->boolean('is_featured')->default(false);

            $table->uuid('author_id');
            $table->uuid('category_id')->nullable();
            $table->uuid('published_by')->nullable();
            $table->uuid('locked_by')->nullable();

            $table->timestamp('published_at')->nullable();
            $table->timestamps();

            $table->foreign('author_id')->references('id')->on('users')->cascadeOnDelete();
            $table->foreign('category_id')->references('id')->on('categories')->nullOnDelete();
            $table->foreign('published_by')->references('id')->on('users')->nullOnDelete();
            $table->foreign('locked_by')->references('id')->on('users')->nullOnDelete();

            $table->index('status');
            $table->index('view_count');
            $table->index('published_at');
        });
    }

    public function down(): void
    {
        Schema::dropIfExists('articles');
    }
};'''

CODE_USER_SEEDER = '''<?php

namespace Database\\Seeders;

use Illuminate\\Database\\Seeder;
use Illuminate\\Support\\Facades\\DB;
use Illuminate\\Support\\Facades\\Hash;

class UserSeeder extends Seeder
{
    public function run(): void
    {
        $users = json_decode(file_get_contents(database_path('data/dummy2.json')), true);

        foreach ($users as $user) {
            DB::table('users')->insert([
                'id'            => $user['id'],
                'name'          => $user['name'],
                'email'         => $user['email'],
                'password'      => Hash::make($user['password']),
                'role'          => $user['role'],
                'is_active'     => $user['is_active'] ?? true,
                'email_verified_at' => now(),
                'created_at'    => now(),
                'updated_at'    => now(),
            ]);
        }
    }
}'''

CODE_SCHEDULED_PUBLISH = '''<?php

namespace App\\Console\\Commands;

use App\\Services\\ScheduledPublishService;
use Illuminate\\Console\\Command;

class PublishScheduledArticles extends Command
{
    protected $signature = 'articles:publish-scheduled';
    protected $description = 'Publish articles that have reached their scheduled time';

    private ScheduledPublishService $service;

    public function __construct(ScheduledPublishService $service)
    {
        parent::__construct();
        $this->service = $service;
    }

    public function handle(): int
    {
        $count = $this->service->publishDueArticles();
        $this->info("Published {$count} scheduled article(s).");
        return Command::SUCCESS;
    }
}'''

CODE_FRONTEND_FOLDER = '''frontend/
├── app/
│   ├── (main)/              # Route group portal publik
│   │   ├── page.tsx         # Dashboard Reader (beranda)
│   │   ├── [slug]/page.tsx  # Detail artikel (SSG)
│   │   └── bookmark/page.tsx
│   ├── cms/                 # Dashboard CMS
│   │   ├── dashboard/page.tsx
│   │   ├── artikel/
│   │   │   ├── page.tsx     # Bank Berita
│   │   │   ├── tulis/page.tsx
│   │   │   └── tinjauan/page.tsx
│   │   ├── media/page.tsx
│   │   ├── komentar/page.tsx
│   │   └── pengguna/page.tsx
│   ├── components/          # Shared components
│   ├── constants/           # roles.ts, errorMessages.ts
│   ├── login/page.tsx
│   ├── register/page.tsx
│   ├── types/index.ts       # TypeScript interfaces
│   ├── layout.tsx           # Root layout
│   └── globals.css
├── lib/
│   ├── api/                 # API functions per module
│   │   ├── articles.ts
│   │   ├── bookmarks.ts
│   │   └── ...
│   ├── auth.ts
│   ├── axios.ts             # Axios instance + interceptors
│   └── validations.ts       # Zod schemas
├── stores/
│   └── authStore.ts         # Zustand auth store
└── utils/
    └── cropImage.ts'''

CODE_DASHBOARD_PAGE = '''"use client";
import { useEffect, useState } from "react";
import { useAuthStore } from "@/stores/authStore";
import axiosInstance from "@/lib/axios";
import {
  BarChart, Bar, XAxis, YAxis, CartesianGrid, Tooltip,
  ResponsiveContainer, PieChart, Pie, Cell, Legend,
} from "recharts";

const COLORS = ["#3B82F6", "#10B981", "#F59E0B", "#EF4444", "#8B5CF6"];

function StatCard({ title, value, color }: { title: string; value: number | string; color: string }) {
  return (
    <div className={`bg-white rounded-lg shadow p-6 border-l-4 ${color}`}>
      <p className="text-sm text-gray-500">{title}</p>
      <p className="text-2xl font-bold mt-2">{value}</p>
    </div>
  );
}

export default function EditorDashboardView() {
  const { user } = useAuthStore();
  const [stats, setStats] = useState<any>(null);

  useEffect(() => {
    axiosInstance.get("/api/cms/statistics").then((res) => {
      setStats(res.data.data);
    });
  }, []);

  if (!stats) return <div>Loading...</div>;

  const statCards = [
    { title: "Artikel Published", value: stats.published, color: "border-green-500" },
    { title: "Draft", value: stats.draft, color: "border-yellow-500" },
    { title: "Menunggu Review", value: stats.review, color: "border-blue-500" },
  ];

  return (
    <div className="p-6 space-y-6">
      <h1 className="text-2xl font-bold">Dashboard — {user?.name}</h1>
      <div className="grid grid-cols-1 md:grid-cols-3 gap-4">
        {statCards.map((s) => (
          <StatCard key={s.title} {...s} />
        ))}
      </div>
      <div className="grid grid-cols-1 lg:grid-cols-2 gap-6">
        <div className="bg-white p-6 rounded-lg shadow">
          <h3 className="font-semibold mb-4">Tren Artikel per Bulan</h3>
          <ResponsiveContainer width="100%" height={300}>
            <BarChart data={stats.yearly_data}>
              <CartesianGrid strokeDasharray="3 3" />
              <XAxis dataKey="month" />
              <YAxis />
              <Tooltip />
              <Bar dataKey="total" fill="#3B82F6" />
            </BarChart>
          </ResponsiveContainer>
        </div>
        <div className="bg-white p-6 rounded-lg shadow">
          <h3 className="font-semibold mb-4">Distribusi Kategori</h3>
          <ResponsiveContainer width="100%" height={300}>
            <PieChart>
              <Pie data={stats.category_data} dataKey="total" nameKey="name" cx="50%" cy="50%" outerRadius={100} label>
                {stats.category_data.map((_: any, i: number) => (
                  <Cell key={i} fill={COLORS[i % COLORS.length]} />
                ))}
              </Pie>
              <Legend />
            </PieChart>
          </ResponsiveContainer>
        </div>
      </div>
    </div>
  );
}'''

CODE_BOOKMARK_API_FRONTEND = '''import axiosInstance from "../axios";
import type { ApiSuccess, Bookmark, Pagination } from "@/app/types";

export async function getBookmarks(page = 1, limit = 10) {
  const res = await axiosInstance.get<ApiSuccess<Bookmark[] & Pagination>>(
    `/api/bookmarks?page=${page}&limit=${limit}`
  );
  return res.data.data;
}

export async function toggleBookmark(articleId: string) {
  const res = await axiosInstance.post<ApiSuccess<{ bookmarked: boolean }>>(
    "/api/bookmarks",
    { article_id: articleId }
  );
  return res.data.data;
}'''

CODE_AXIOS_INTERCEPTOR = '''import axios from "axios";
import { getAccessToken, setAccessToken, useAuthStore } from "@/stores/authStore";

const axiosInstance = axios.create({
  baseURL: process.env.NEXT_PUBLIC_API_URL || "http://localhost:8000",
});

// Request interceptor: attach Bearer token
axiosInstance.interceptors.request.use((config) => {
  const token = getAccessToken();
  if (token) {
    config.headers.Authorization = `Bearer ${token}`;
  }
  return config;
});

// Response interceptor: silent token refresh
axiosInstance.interceptors.response.use(
  (response) => response,
  async (error) => {
    const originalRequest = error.config;
    const errorCode = error.response?.data?.error;

    if (errorCode === "TOKEN_EXPIRED" && !originalRequest._retry) {
      originalRequest._retry = true;
      try {
        const refreshToken = localStorage.getItem("refresh_token");
        const res = await axios.post(
          `${process.env.NEXT_PUBLIC_API_URL}/api/auth/refresh`,
          { refresh_token: refreshToken }
        );
        const newToken = res.data.data.access_token;
        setAccessToken(newToken);
        originalRequest.headers.Authorization = `Bearer ${newToken}`;
        return axiosInstance(originalRequest);
      } catch {
        useAuthStore.getState().logout();
        window.location.href = "/login";
        return Promise.reject(error);
      }
    }
    return Promise.reject(error);
  }
);

export default axiosInstance;'''

CODE_TYPESCRIPT_INTERFACES = '''export interface User {
  id: string;
  name: string;
  email: string;
  role: "reader" | "journalist" | "editor" | "admin";
  is_active: boolean;
  email_verified_at: string | null;
  created_at: string;
}

export interface Category {
  id: string;
  name: string;
  slug: string;
  description: string | null;
}

export interface Tag {
  id: string;
  name: string;
  slug: string;
}

export interface Media {
  id: string;
  filename: string;
  url: string;
  mime_type: string;
  size: number;
  alt_text: string | null;
  created_at: string;
}

export interface Article {
  id: string;
  title: string;
  slug: string;
  excerpt: string | null;
  content: string;
  cover_url: string | null;
  status: "draft" | "review" | "published" | "scheduled" | "rejected" | "archived";
  view_count: number;
  is_featured: boolean;
  author: Pick<User, "id" | "name">;
  category: Pick<Category, "id" | "name" | "slug"> | null;
  tags: Pick<Tag, "id" | "name" | "slug">[];
  published_at: string | null;
  created_at: string;
}

export interface Comment {
  id: string;
  content: string;
  status: "pending" | "approved" | "rejected";
  user: Pick<User, "id" | "name">;
  parent_id: string | null;
  replies?: Comment[];
  created_at: string;
}

export interface Pagination {
  total: number;
  page: number;
  limit: number;
  total_pages: number;
}

export interface ApiSuccess<T> {
  status: "success";
  data: T;
}

export interface ApiError {
  status: "error";
  error: string;
  message: string;
  details?: Record<string, string[]>;
}'''

CODE_ZOD_SCHEMAS = '''import { z } from "zod";

export const loginSchema = z.object({
  email: z.string().email("Email tidak valid"),
  password: z.string().min(6, "Password minimal 6 karakter"),
});

export const registerSchema = z.object({
  name: z.string().min(3, "Nama minimal 3 karakter").max(100),
  email: z.string().email("Email tidak valid"),
  password: z.string().min(6, "Password minimal 6 karakter"),
  password_confirmation: z.string(),
}).refine((data) => data.password === data.password_confirmation, {
  message: "Konfirmasi password tidak cocok",
  path: ["password_confirmation"],
});

export const articleSchema = z.object({
  title: z.string().min(5, "Judul minimal 5 karakter").max(255),
  slug: z.string().optional(),
  excerpt: z.string().max(500).optional(),
  content: z.string().min(10, "Konten minimal 10 karakter"),
  category_id: z.string().uuid("Kategori wajib dipilih"),
  tag_ids: z.array(z.string().uuid()).default([]),
  cover_url: z.string().nullable().optional(),
  is_featured: z.boolean().default(false),
});

export const scheduleSchema = z.object({
  scheduled_at: z.string().refine((val) => {
    const scheduledTime = new Date(val);
    const now = new Date();
    const diff = (scheduledTime.getTime() - now.getTime()) / 60000;
    return diff >= 5;
  }, "Waktu tayang minimal 5 menit dari sekarang"),
});

export const commentSchema = z.object({
  content: z.string().min(1, "Komentar tidak boleh kosong").max(1000),
  parent_id: z.string().uuid().nullable().optional(),
});

export const createUserSchema = z.object({
  name: z.string().min(3).max(100),
  email: z.string().email("Email tidak valid"),
  password: z.string().min(6, "Password minimal 6 karakter"),
  role: z.enum(["reader", "journalist", "editor", "admin"]),
  is_active: z.boolean().default(true),
});

export const profileSchema = z.object({
  name: z.string().min(3).max(100),
  email: z.string().email("Email tidak valid"),
});

export const changePasswordSchema = z.object({
  current_password: z.string().min(1, "Password lama wajib diisi"),
  new_password: z.string().min(6, "Password baru minimal 6 karakter"),
  new_password_confirmation: z.string(),
}).refine((data) => data.new_password === data.new_password_confirmation, {
  message: "Konfirmasi password baru tidak cocok",
  path: ["new_password_confirmation"],
});'''

CODE_ROLES_CONSTANTS = '''// roles.ts
export const ROLES = {
  READER: "reader",
  JOURNALIST: "journalist",
  EDITOR: "editor",
  ADMIN: "admin",
} as const;

export function canPublish(role: string): boolean {
  return role === ROLES.EDITOR || role === ROLES.ADMIN;
}

export function canDeleteComment(role: string): boolean {
  return role === ROLES.EDITOR || role === ROLES.ADMIN;
}

export function canManageUsers(role: string): boolean {
  return role === ROLES.ADMIN;
}'''

CODE_ARTICLES_API_FRONTEND = '''import axiosInstance from "../axios";
import type { ApiSuccess, Article, Comment, Pagination } from "@/app/types";

export async function getArticles(page = 1, limit = 10) {
  const res = await axiosInstance.get<ApiSuccess<Article[] & Pagination>>(
    `/api/articles?page=${page}&limit=${limit}`
  );
  return res.data.data;
}

export async function getCmsArticles(page = 1, limit = 10, filters?: Record<string, string>) {
  const params = new URLSearchParams({ page: String(page), limit: String(limit) });
  if (filters) Object.entries(filters).forEach(([k, v]) => params.append(k, v));
  const res = await axiosInstance.get<ApiSuccess<Article[] & Pagination>>(
    `/api/cms/articles?${params}`
  );
  return res.data.data;
}

export async function getArticleBySlug(slug: string) {
  const res = await axiosInstance.get<ApiSuccess<Article>>(`/api/articles/${slug}`);
  return res.data.data;
}

export async function createArticle(data: FormData) {
  const res = await axiosInstance.post<ApiSuccess<Article>>("/api/cms/articles", data, {
    headers: { "Content-Type": "multipart/form-data" },
  });
  return res.data.data;
}

export async function updateArticleStatus(id: string, status: string, changeNote?: string) {
  const res = await axiosInstance.patch<ApiSuccess<Article>>(`/api/cms/articles/${id}/status`, {
    status, change_note: changeNote,
  });
  return res.data.data;
}

export async function updateArticle(id: string, data: FormData) {
  const res = await axiosInstance.post<ApiSuccess<Article>>(`/api/cms/articles/${id}?_method=PUT`, data, {
    headers: { "Content-Type": "multipart/form-data" },
  });
  return res.data.data;
}

export async function deleteArticle(id: string) {
  await axiosInstance.delete(`/api/cms/articles/${id}`);
}

export async function lockArticle(id: string) {
  await axiosInstance.post(`/api/cms/articles/${id}/lock`);
}

export async function unlockArticle(id: string) {
  await axiosInstance.post(`/api/cms/articles/${id}/unlock`);
}

export async function getComments(articleId: string) {
  const res = await axiosInstance.get<ApiSuccess<Comment[]>>(`/api/articles/${articleId}/comments`);
  return res.data.data;
}

export async function postComment(articleId: string, content: string, parentId?: string | null) {
  const res = await axiosInstance.post<ApiSuccess<Comment>>(`/api/articles/${articleId}/comments`, {
    content, parent_id: parentId,
  });
  return res.data.data;
}

export async function deleteComment(id: string) {
  await axiosInstance.delete(`/api/comments/${id}`);
}'''

CODE_AUTH_STORE = '''import { create } from "zustand";
import { persist } from "zustand/middleware";
import type { User } from "@/app/types";

interface AuthState {
  accessToken: string | null;
  user: User | null;
  isAuthenticated: boolean;
  setAuth: (token: string, user: User) => void;
  setAccessToken: (token: string) => void;
  updateUser: (user: User) => void;
  logout: () => void;
}

export const useAuthStore = create<AuthState>()(
  persist(
    (set) => ({
      accessToken: null,
      user: null,
      isAuthenticated: false,

      setAuth: (token, user) =>
        set({ accessToken: token, user, isAuthenticated: true }),

      setAccessToken: (token) => set({ accessToken: token }),

      updateUser: (user) => set({ user }),

      logout: () => {
        localStorage.removeItem("refresh_token");
        set({ accessToken: null, user: null, isAuthenticated: false });
      },
    }),
    { name: "auth-storage" }
  )
);

// Helper functions for Axios interceptor
export const getAccessToken = () => useAuthStore.getState().accessToken;
export const setAccessToken = (token: string) =>
  useAuthStore.getState().setAccessToken(token);'''

# ============================================================
# API EXAMPLE CODE SNIPPETS
# ============================================================

CODE_API_LOGIN = '''// POST /api/auth/login
// Request:
{
  "email": "jurnalis@ketik.com",
  "password": "password123"
}

// Response (200 OK):
{
  "status": "success",
  "data": {
    "access_token": "eyJ0eXAiOiJKV1QiLCJhbGc...",
    "refresh_token": "a1b2c3d4e5f6...",
    "token_type": "bearer",
    "expires_in": 900,
    "user": {
      "id": "uuid-here",
      "name": "Jurnalis Klojen",
      "email": "jurnalis@ketik.com",
      "role": "journalist"
    }
  }
}'''

CODE_API_DETAIL_ARTICLE = '''// GET /api/articles/{slug}
// Response (200 OK):
{
  "status": "success",
  "data": {
    "id": "uuid-here",
    "title": "Memulai Belajar Next.js 15",
    "slug": "memulai-belajar-nextjs-15",
    "excerpt": "Panduan lengkap untuk memulai...",
    "content": "<p>Next.js 15 adalah framework...</p>",
    "cover_url": "https://storage.ketik.com/images/cover.jpg",
    "status": "published",
    "view_count": 1523,
    "is_featured": true,
    "author": { "id": "uuid", "name": "Jurnalis Klojen" },
    "category": { "id": "uuid", "name": "Teknologi", "slug": "teknologi" },
    "tags": [
      { "id": "uuid", "name": "Next.js", "slug": "nextjs" },
      { "id": "uuid", "name": "React", "slug": "react" }
    ],
    "published_at": "2026-05-01T10:00:00+07:00",
    "created_at": "2026-04-28T08:30:00+07:00"
  }
}'''

CODE_API_LIST_ARTICLES = '''// GET /api/articles?page=1&limit=10
// Response (200 OK):
{
  "status": "success",
  "data": {
    "data": [
      {
        "id": "uuid",
        "title": "Memulai Belajar Next.js 15",
        "slug": "memulai-belajar-nextjs-15",
        "excerpt": "Panduan lengkap...",
        "cover_url": "https://storage.ketik.com/images/cover.jpg",
        "category": { "name": "Teknologi" },
        "author": { "name": "Jurnalis Klojen" },
        "published_at": "2026-05-01T10:00:00+07:00",
        "view_count": 1523
      }
    ],
    "total": 45,
    "page": 1,
    "limit": 10,
    "total_pages": 5
  }
}'''

CODE_API_TOGGLE_BOOKMARK = '''// POST /api/bookmarks
// Request:
{
  "article_id": "uuid-article-here"
}

// Response (200 OK) — bookmark added:
{
  "status": "success",
  "data": { "bookmarked": true }
}

// Response (200 OK) — bookmark removed:
{
  "status": "success",
  "data": { "bookmarked": false }
}'''

CODE_API_CHANGE_STATUS = '''// PATCH /api/cms/articles/{id}/status
// Request:
{
  "status": "published",
  "change_note": "Artikel telah direview dan disetujui"
}

// Response (200 OK):
{
  "status": "success",
  "data": {
    "id": "uuid",
    "title": "Memulai Belajar Next.js 15",
    "status": "published",
    "published_at": "2026-06-01T14:30:00+07:00",
    "published_by": { "id": "uuid", "name": "Editor Klojen" }
  }
}'''

CODE_API_SEARCH = '''// GET /api/search?page=1&q=nextjs
// Response (200 OK):
{
  "status": "success",
  "data": {
    "data": [
      {
        "id": "uuid",
        "title": "Memulai Belajar Next.js 15",
        "slug": "memulai-belajar-nextjs-15",
        "excerpt": "Panduan lengkap...",
        "cover_url": "https://storage.ketik.com/images/cover.jpg",
        "category": { "name": "Teknologi" },
        "published_at": "2026-05-01T10:00:00+07:00"
      }
    ],
    "total": 3,
    "page": 1,
    "limit": 10,
    "total_pages": 1
  }
}'''

CODE_API_UPLOAD_MEDIA = '''// POST /api/media
// Request: multipart/form-data
//   file: (binary image file)
//   alt_text: "Deskripsi gambar"

// Response (201 Created):
{
  "status": "success",
  "data": {
    "id": "uuid-media",
    "filename": "banner-nextjs.jpg",
    "url": "https://storage.ketik.com/media/banner-nextjs.jpg",
    "mime_type": "image/jpeg",
    "size": 245678,
    "alt_text": "Deskripsi gambar",
    "created_at": "2026-06-01T14:30:00+07:00"
  }
}'''

# ============================================================
# DIAGRAM CONTENT
# ============================================================

DIAGRAM_ARCHITECTURE = '''[Client Layer — Browser]
        |
        v
[Frontend — Next.js 15 (SSR/SSG/CSR)]
   ├── Portal Publik (Beranda, Detail, Bookmark)
   └── CMS Redaksi (Dashboard, Artikel, Media, Komentar)
        |
        |  HTTP REST API (JSON)
        |  Authorization: Bearer <JWT>
        v
[API Gateway — Laravel 12 Middleware]
   ├── AuthenticateJwt (token verification)
   ├── EnsureAdmin (role authorization)
   └── ThrottleRequests (rate limiting)
        |
        v
[Controller Layer — Thin Controllers]
   ├── AuthController, ArticleController, BookmarkController
   ├── CmsArticleController, CmsCommentController, MediaController
   └── UserController, AnalyticsController
        |
        v
[Service Layer — Business Logic]
   ├── AuthService, CmsArticleService, BookmarkService
   ├── SearchService, ScheduledPublishService, MediaService
   └── CommentService, UserService
        |
        v
[Repository Layer — Data Access]
   ├── ArticleRepository, UserRepository, MediaRepository
   ├── CommentRepository, CategoryRepository, TagRepository
   └── BookmarkRepository, SearchIndexRepository
        |
        v
[Database — MySQL 8]
   ├── 14 Tabel (users, articles, categories, tags, comments, ...)
   ├── FULLTEXT Index (search_indexes)
   └── UUID Primary Keys'''

DIAGRAM_DESIGN_PATTERNS = '''Service-Repository Pattern:

  [HTTP Request]
       |
       v
  [Controller] ── thin, hanya terima request & return response
       |
       |  delegate
       v
  [Service] ── business logic, validasi role, orkestrasi
       |
       |  call interface
       v
  [Repository] ── query database (Eloquent ORM / Query Builder)
       |
       v
  [Database]

Dependency Injection via RepositoryServiceProvider:
  ArticleRepositoryInterface     → ArticleRepository
  UserRepositoryInterface        → UserRepository
  CommentRepositoryInterface     → CommentRepository
  MediaRepositoryInterface       → MediaRepository
  CategoryRepositoryInterface    → CategoryRepository
  TagRepositoryInterface         → TagRepository
  BookmarkRepositoryInterface    → BookmarkRepository
  SearchIndexRepositoryInterface → SearchIndexRepository'''

DIAGRAM_BACKEND_FOLDER = '''backend/
├── app/
│   ├── Console/Commands/
│   │   ├── PublishScheduledArticles.php
│   │   └── ReindexArticlesCommand.php
│   ├── Http/
│   │   ├── Controllers/    (16 controllers)
│   │   │   ├── AuthController.php
│   │   │   ├── ArticleController.php
│   │   │   ├── BookmarkController.php
│   │   │   ├── CmsArticleController.php
│   │   │   ├── CmsCommentController.php
│   │   │   ├── CmsDashboardController.php
│   │   │   ├── MediaController.php
│   │   │   ├── UserController.php
│   │   │   └── ...
│   │   └── Middleware/
│   │       └── EnsureAdmin.php
│   ├── Mail/
│   │   ├── NewUserCredentials.php
│   │   └── ResetPasswordMail.php
│   ├── Models/             (12 models)
│   │   ├── Article.php, User.php, Category.php, Tag.php
│   │   ├── Comment.php, Media.php, Bookmark.php
│   │   └── ...
│   ├── Providers/
│   │   ├── AppServiceProvider.php
│   │   └── RepositoryServiceProvider.php
│   ├── Repositories/
│   │   ├── Contracts/      (9 interfaces)
│   │   ├── ArticleRepository.php
│   │   ├── BookmarkRepository.php
│   │   └── ...
│   └── Services/           (12 services)
│       ├── AuthService.php
│       ├── BookmarkService.php
│       ├── CmsArticleService.php
│       ├── ScheduledPublishService.php
│       └── ...
├── database/
│   ├── migrations/         (26 migration files)
│   ├── seeders/            (8 seeder files)
│   └── data/dummy2.json
├── routes/
│   ├── api.php             (132 lines, 50+ endpoints)
│   └── console.php
└── config/
    ├── jwt.php, cors.php, database.php, mail.php
    └── ...'''

# ============================================================
# FILLING LOGIC
# ============================================================

def fill_code_placeholder(paragraphs, idx, code_text):
    """Replace 'KODE' paragraph with actual code lines."""
    lines = code_text.strip().split('\n')
    p = paragraphs[idx]
    # Clear existing text
    for run in p.runs:
        run.text = ""
    # Set first line
    if p.runs:
        p.runs[0].text = lines[0]
    else:
        run = p.add_run(lines[0])
        run.font.name = "Courier New"
        run.font.size = Pt(8)

    # Insert remaining lines as new paragraphs after current
    insert_point = p._element
    for line in lines[1:]:
        new_p = docx.oxml.OxmlElement('w:p')
        new_r = docx.oxml.OxmlElement('w:r')
        new_rpr = docx.oxml.OxmlElement('w:rPr')
        new_rfont = docx.oxml.OxmlElement('w:rFonts')
        new_rfont.set(qn('w:ascii'), 'Courier New')
        new_rfont.set(qn('w:hAnsi'), 'Courier New')
        new_rpr.append(new_rfont)
        new_sz = docx.oxml.OxmlElement('w:sz')
        new_sz.set(qn('w:val'), '16')
        new_rpr.append(new_sz)
        new_r.append(new_rpr)
        new_t = docx.oxml.OxmlElement('w:t')
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = line
        new_r.append(new_t)
        new_p.append(new_r)
        insert_point.addnext(new_p)
        insert_point = new_p

def fill_gambar_placeholder(paragraphs, idx, diagram_text):
    """Replace 'GAMBAR' paragraph with diagram content."""
    fill_code_placeholder(paragraphs, idx, diagram_text)

paragraphs = doc.paragraphs
filled_count = 0

# Track which KODE/GAMBAR we've filled (by order of appearance)
kode_index = 0
gambar_index = 0

# Define the order of KODE placeholders based on document analysis
KODE_ORDER = [
    CODE_BOOKMARK_SERVICE_TOGGLE,   # already filled as Potongan Kode 1
    CODE_BOOKMARK_SERVICE_TOGGLE,   # getBookmarks method (KODE after line 563)
    CODE_BOOKMARK_CONTROLLER,       # KODE after line 566
    CODE_CMS_DASHBOARD_INDEX,       # KODE after line 574
    CODE_EDITOR_STATS,              # KODE after line 577
    CODE_API_ROUTES,                # KODE after line 581
    CODE_API_LOGIN,                 # KODE after line 599
    CODE_API_DETAIL_ARTICLE,        # KODE after line 601
    CODE_API_LIST_ARTICLES,         # KODE after line 603
    CODE_API_TOGGLE_BOOKMARK,       # KODE after line 605
    CODE_API_CHANGE_STATUS,         # KODE after line 607
    CODE_API_SEARCH,                # KODE after line 609
    CODE_API_UPLOAD_MEDIA,          # KODE after line 611
    CODE_MIDDLEWARE_ADMIN,          # KODE after line 614
    CODE_ARTICLES_MIGRATION,        # KODE after line 618
    CODE_USER_SEEDER,               # KODE after line 621
    CODE_SCHEDULED_PUBLISH,         # KODE after line 625
    CODE_FRONTEND_FOLDER,           # KODE after line 633
    CODE_DASHBOARD_PAGE,            # KODE after line 639
    CODE_DASHBOARD_PAGE,            # KODE after line 642 (EditorDashboardView)
    CODE_BOOKMARK_API_FRONTEND,     # KODE after line 677
    CODE_AXIOS_INTERCEPTOR,         # KODE after line 682
    CODE_TYPESCRIPT_INTERFACES,     # KODE after line 711
    CODE_ZOD_SCHEMAS,               # KODE after line 715
    CODE_ROLES_CONSTANTS,           # KODE after line 727
    CODE_ARTICLES_API_FRONTEND,     # KODE after line 730
    CODE_AUTH_STORE,                # KODE after line 734 (authStore)
]

# Define GAMBAR placeholders order
GAMBAR_ORDER = [
    DIAGRAM_ARCHITECTURE,    # Arsitektur Sistem
    DIAGRAM_DESIGN_PATTERNS, # Design Patterns
    DIAGRAM_BACKEND_FOLDER,  # Struktur Folder Backend
]

# Process paragraphs
i = 0
kode_filled = 0
gambar_filled = 0

while i < len(paragraphs):
    text = paragraphs[i].text.strip()

    if text == "KODE" and kode_filled < len(KODE_ORDER):
        fill_code_placeholder(paragraphs, i, KODE_ORDER[kode_filled])
        kode_filled += 1
        filled_count += 1
        print(f"  Filled KODE #{kode_filled} at paragraph {i}")

    elif text == "GAMBAR" and gambar_filled < len(GAMBAR_ORDER):
        fill_gambar_placeholder(paragraphs, i, GAMBAR_ORDER[gambar_filled])
        gambar_filled += 1
        filled_count += 1
        print(f"  Filled GAMBAR #{gambar_filled} at paragraph {i}")

    i += 1

# Save
output_path = "23081010208_laporan_pkl_filled.docx"
doc.save(output_path)
print(f"\nDone! Filled {filled_count} placeholders ({kode_filled} KODE, {gambar_filled} GAMBAR)")
print(f"Saved to: {output_path}")
