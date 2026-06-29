"""Generate logbook.docx matching laporan_plk_2 content."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT = r'c:\apps\klojen.com\logbook3.docx'

logs = [
    ('Minggu 1-3 / 9 Feb – 28 Feb 2026',
     'Pengenalan perusahaan, pelatihan jurnalistik, penulisan dan publikasi berita, serta pemahaman alur kerja redaksi.',
     'Memahami alur kerja dan pembagian tugas selama magang. Memperoleh pengetahuan dasar mengenai penulisan berita dan teknik pengambilan foto jurnalistik.'),
    ('Minggu 4-8 / 2 Mar – 4 Apr 2026',
     'Pengenalan proyek, Git dan GitHub, pengelolaan data website, serta pembelajaran workflow pengembangan sistem.',
     'Memahami dasar penggunaan Git dan GitHub serta alur kerja kolaborasi proyek.'),
    ('Minggu 9 / 6-11 Apr 2026',
     'Merancang UI/UX menggunakan Figma untuk halaman CMS (artikel, tulis berita, tinjauan, profil) dan portal berita (beranda, detail, kategori).',
     'Berhasil membuat desain UI/UX yang responsif dan modern menggunakan Figma.'),
    ('Minggu 10 / 13-18 Apr 2026',
     'Melakukan perancangan sistem menggunakan Activity Diagram serta perancangan database.',
     'Berhasil menyusun dokumentasi perancangan sistem.'),
    ('Minggu 11 / 20-25 Apr 2026',
     'Setup project Next.js 15 dan Tailwind CSS. Mulai implementasi layout dasar dan halaman-halaman CMS.',
     'Project frontend berhasil di-setup dengan layout dasar yang berfungsi.'),
    ('Minggu 12 / 27 Apr – 2 Mei 2026',
     'Pengembangan halaman daftar artikel, tulis berita, dan profil berdasarkan desain Figma.',
     'Halaman artikel dan profil berhasil dikembangkan sesuai desain.'),
    ('Minggu 13 / 4-9 Mei 2026',
     'Melanjutkan pengembangan halaman tinjauan artikel, preview berita, dan rich text editor untuk tulis berita.',
     'Fitur rich text editor, tinjauan artikel, dan preview berita berhasil dikembangkan.'),
    ('Minggu 14 / 11-16 Mei 2026',
     'Pengembangan frontend portal berita: beranda, detail artikel, dan halaman kategori.',
     'Portal berita berhasil dikembangkan dengan tampilan yang responsif.'),
    ('Minggu 15 / 18-23 Mei 2026',
     'Mulai integrasi frontend dengan REST API backend. Implementasi fitur upload gambar dan crop avatar.',
     'Frontend berhasil terhubung dengan API dan fitur upload berjalan.'),
    ('Minggu 16 / 25-30 Mei 2026',
     'Pengembangan backend: endpoint create article, edit profile, dan change password menggunakan Laravel 12.',
     'Backend berhasil dibuat dan terintegrasi dengan frontend.'),
    ('Minggu 17 / 1-6 Jun 2026',
     'Penyempurnaan fitur: memperbaiki bug, optimasi komponen, dan memperbaiki struktur kode.',
     'Beberapa bug berhasil diperbaiki dan performa menjadi lebih baik.'),
    ('Minggu 18 / 8-13 Jun 2026',
     'Penyusunan Laporan Akhir Magang Mandiri (PLK): menulis Bab I (Pendahuluan), Bab II (Organisasi Mitra PLK), Bab III (Rancang Bangun Sistem), dan Bab IV (Rekognisi Mata Kuliah). Menyusun dokumentasi perancangan sistem (Activity Diagram) serta potongan kode dan penjelasan untuk setiap fitur yang dikembangkan.',
     'Draft laporan PLK Bab I-IV berhasil disusun lengkap dengan dokumentasi perancangan sistem dan rekognisi mata kuliah (Analisis Kebutuhan, Pemrograman API, Uji Coba dan Implementasi).'),
    ('Minggu 19 / 15-19 Jun 2026',
     'Penyusunan Laporan PKL, luaran jurnal ilmiah, serta finalisasi seluruh dokumen. Finalisasi project: pengecekan keseluruhan fitur, penyempurnaan laporan PLK dan PKL, penyusunan jurnal, serta persiapan presentasi hasil magang.',
     'Laporan PLK, Laporan PKL, dan luaran jurnal berhasil diselesaikan. Seluruh dokumen dan bahan presentasi berhasil disusun untuk laporan akhir magang.'),
]

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Lampiran A. Log Activity')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = sub.add_run('Berikut adalah log activity kegiatan selama mengikuti Program Magang Mandiri di PT. Ketik Media Siber:')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

for row in table.rows:
    row.cells[0].width = Cm(3.0)
    row.cells[1].width = Cm(6.5)
    row.cells[2].width = Cm(5.5)
    row.cells[3].width = Cm(2.5)

headers = ['Minggu/Tanggal', 'Kegiatan', 'Hasil', 'Validasi\nDosen Pembimbing']
hdr_row = table.rows[0]
for i, text in enumerate(headers):
    cell = hdr_row.cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'E8EDF3',
    })
    shading.append(shading_elem)

for minggu, kegiatan, hasil in logs:
    row = table.add_row()
    
    p1 = row.cells[0].paragraphs[0]
    run1 = p1.add_run(minggu)
    run1.bold = True
    run1.font.size = Pt(10)
    run1.font.name = 'Times New Roman'
    
    p2 = row.cells[1].paragraphs[0]
    run2 = p2.add_run(kegiatan)
    run2.font.size = Pt(10)
    run2.font.name = 'Times New Roman'
    
    p3 = row.cells[2].paragraphs[0]
    run3 = p3.add_run(hasil)
    run3.font.size = Pt(10)
    run3.font.name = 'Times New Roman'
    
    p4 = row.cells[3].paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

for row in table.rows:
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = tcPr.makeelement(qn('w:vAlign'), {qn('w:val'): 'top'})
        tcPr.append(vAlign)

doc.save(OUTPUT)
print(f'Done! Saved to {OUTPUT}')
