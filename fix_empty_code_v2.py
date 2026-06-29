#!/usr/bin/env python3
"""
Script to fill empty code snippets and diagram placeholders in FIX.docx.
"""
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_code_format(paragraph, text):
    """Format paragraph as code (monospace, small font)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ''
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_text_after_paragraph(paragraph, text, is_code=False):
    """Add a new paragraph after the given paragraph."""
    new_para = docx.oxml.OxmlElement('w:p')
    paragraph._p.addnext(new_para)
    new_para_obj = docx.text.paragraph.Paragraph(new_para, paragraph._parent)
    if is_code:
        set_code_format(new_para_obj, text)
    else:
        new_para_obj.add_run(text)
    return new_para_obj

def fill_code_snippets_v2(doc):
    """Fill all empty code snippets in the document with better logic."""
    
    # Map of code snippet labels to their actual code
    code_snippets = {
        "Potongan Kode — Format respons error API": '''{
  "status": "error",
  "error": "ERROR_CODE",
  "message": "Pesan error dalam bahasa Indonesia",
  "details": { ... }
}''',
        
        "Potongan Kode 1. Method toggle()": '''public function toggle(int $userId, string $articleId): bool
{
    $existing = DB::table('bookmarks')
        ->where('user_id', $userId)
        ->where('article_id', $articleId)
        ->first();

    if ($existing) {
        DB::table('bookmarks')
            ->where('user_id', $userId)
            ->where('article_id', $articleId)
            ->delete();
        return false;
    }

    DB::table('bookmarks')->insert([
        'id'         => Str::uuid()->toString(),
        'user_id'    => $userId,
        'article_id' => $articleId,
        'created_at' => now(),
    ]);
    return true;
}''',
        
        "Potongan Kode 2. Method getBookmarks()": '''public function getBookmarks(int $userId): array
{
    $rows = DB::table('bookmarks')
        ->where('user_id', $userId)
        ->orderByDesc('created_at')
        ->get();

    if ($rows->isEmpty()) return [];

    $articleIds = $rows->pluck('article_id')->toArray();
    $articles = Article::with(['category', 'author'])
        ->whereIn('id', $articleIds)
        ->get()
        ->keyBy('id');

    return $rows->map(function ($row) use ($articles): ?array {
        $article = $articles->get($row->article_id);
        if (!$article) return null;
        return [
            'id' => $row->id,
            'created_at' => $row->created_at,
            'article' => [
                'id' => $article->id,
                'title' => $article->title,
                'slug' => $article->slug,
                'featured_image_url' => $article->featured_image_url,
                'category' => $article->category ? [
                    'id' => $article->category->id,
                    'name' => $article->category->name,
                ] : null,
                'author' => $article->author ? [
                    'id' => $article->author->id,
                    'name' => $article->author->name,
                ] : null,
            ],
        ];
    })->filter()->values()->toArray();
}''',
        
        "Potongan Kode 3. BookmarkController::toggle()": '''public function toggle(Request $request): JsonResponse
{
    $validated = $request->validate([
        'article_id' => 'required|string',
    ]);

    $userId = auth('api')->id();
    $result = $this->bookmarkService->toggle($userId, $validated['article_id']);

    return response()->json([
        'status' => 'success',
        'data'   => ['bookmarked' => $result],
    ]);
}''',
        
        "Potongan Kode 4. CmsDashboardController::index()": '''public function index(Request $request)
{
    $user = $request->user();
    $role = $user->role;

    if ($role === 'admin') {
        return $this->getAdminStats();
    } else {
        return $this->getEditorStats($user);
    }
}''',
        
        "Potongan Kode 5. Method getEditorStats()": '''private function getEditorStats($user)
{
    $baseQuery = function($q) use ($user) {
        if ($user->role === 'journalist') {
            $q->where('author_id', $user->id);
        } elseif ($user->role === 'editor') {
            $q->where(function ($q1) use ($user) {
                $q1->where('author_id', $user->id)
                   ->orWhere('published_by', $user->id);
            });
        }
    };

    $beritaPublish = Article::where($baseQuery)->where('status', 'published')->count();
    $draft = Article::where($baseQuery)->where('status', 'draft')->count();
    $kategoriAktif = Article::where($baseQuery)->distinct('category_id')->count('category_id');

    $yearlyData = Article::where($baseQuery)
        ->select(DB::raw('YEAR(created_at) as year'), DB::raw('count(*) as berita'))
        ->groupBy('year')->orderBy('year', 'asc')->get();

    $categoryData = Article::where($baseQuery)
        ->join('categories', 'articles.category_id', '=', 'categories.id')
        ->select('categories.name', DB::raw('count(*) as value'))
        ->groupBy('categories.id', 'categories.name')->get();

    return response()->json([
        'statCards' => compact('beritaPublish', 'draft', 'kategoriAktif'),
        'yearlyData' => $yearlyData,
        'categoryData' => $categoryData,
    ]);
}''',
        
        "Potongan Kode 6. Definisi route API": '''// routes/api.php

// Bookmarks (requires authentication)
Route::middleware('auth:api')->group(function () {
    Route::get('/bookmarks', [BookmarkController::class, 'index']);
    Route::post('/bookmarks', [BookmarkController::class, 'toggle']);
});

// CMS (requires authentication + role check)
Route::middleware('auth:api')->prefix('cms')->group(function () {
    Route::get('/statistics', [CmsDashboardController::class, 'index']);
    Route::get('/articles', [CmsArticleController::class, 'index']);
    Route::post('/articles', [CmsArticleController::class, 'store']);
    Route::patch('/articles/{id}/status', [CmsArticleController::class, 'updateStatus']);
});''',
        
        "Potongan Kode 7. Komponen DashboardPage": '''// app/cms/dashboard/page.tsx
'use client';
import { useAuthStore } from '@/stores/authStore';
import AdminDashboard from '@/app/components/cms/AdminDashboard';

export default function DashboardPage() {
  const { user } = useAuthStore();
  
  if (user?.role === 'admin') {
    return <AdminDashboard />;
  }
  return <EditorDashboardView />;
}''',
        
        "Potongan Kode 8. EditorDashboardView": '''function EditorDashboardView() {
  const [yearlyData, setYearlyData] = useState<any[]>([]);
  const [categoryData, setCategoryData] = useState<any[]>([]);
  const [statCards, setStatCards] = useState({
    beritaPublish: 0, draft: 0, kategoriAktif: 0
  });

  useEffect(() => {
    axiosInstance.get('/cms/statistics')
      .then(res => {
        setYearlyData(res.data.yearlyData || []);
        setCategoryData(res.data.categoryData || []);
        setStatCards(res.data.statCards || {});
      });
  }, []);

  return (
    <div>
      <StatCard label="Berita Publish" value={statCards.beritaPublish} />
      <StatCard label="Draft" value={statCards.draft} />
      <BarChart data={yearlyData}>
        <Bar dataKey="berita" />
      </BarChart>
      <PieChart data={categoryData} />
    </div>
  );
}''',
        
        "Potongan Kode 9. Fungsi API frontend untuk bookmark": '''// lib/api/bookmarks.ts
import axiosInstance from '@/lib/axios';
import { ApiSuccess, Bookmark } from '@/app/types';

export const getBookmarks = () =>
  axiosInstance.get<ApiSuccess<{ bookmarks: Bookmark[] }>>('/bookmarks');

export const toggleBookmark = (articleId: string) =>
  axiosInstance.post<ApiSuccess<{ bookmarked: boolean }>>(
    '/bookmarks', 
    { article_id: articleId }
  );''',
        
        "Potongan Kode 10. Axios Interceptor": '''// lib/axios.ts
import axios from 'axios';
import { getAccessToken, setAccessToken, logoutStore } from '@/stores/authStore';

const axiosInstance = axios.create({
  baseURL: process.env.NEXT_PUBLIC_API_URL,
});

axiosInstance.interceptors.request.use((config) => {
  const token = getAccessToken();
  if (token) {
    config.headers.Authorization = `Bearer ${token}`;
  }
  return config;
});

axiosInstance.interceptors.response.use(
  (response) => response,
  async (error) => {
    const original = error.config;
    if (error.response?.data?.error === 'TOKEN_EXPIRED' && !original._retry) {
      original._retry = true;
      const res = await axiosInstance.post('/auth/refresh', {
        refresh_token: getRefreshToken()
      });
      const newToken = res.data.data.access_token;
      setAccessToken(newToken);
      original.headers.Authorization = `Bearer ${newToken}`;
      return axiosInstance(original);
    }
    return Promise.reject(error);
  }
);''',
        
        "Potongan Kode 11. Middleware EnsureAdmin": '''// app/Http/Middleware/EnsureAdmin.php
class EnsureAdmin
{
    public function handle(Request $request, Closure $next): Response
    {
        $user = auth('api')->user();

        if (! $user || $user->role !== 'admin') {
            return response()->json([
                'status'  => 'error',
                'code'    => 403,
                'error'   => 'FORBIDDEN',
                'message' => 'Akses ditolak. Hanya admin yang dapat mengakses endpoint ini.',
            ], 403);
        }

        return $next($request);
    }
}''',
        
        "Potongan Kode 12. Migration tabel articles": '''// database/migrations/2026_05_19_100004_create_articles_table.php
return new class extends Migration
{
    public function up(): void
    {
        Schema::create('articles', function (Blueprint $table) {
            $table->uuid('id')->primary();
            $table->foreignId('author_id')->constrained('users')->cascadeOnDelete();
            $table->foreignId('published_by')->nullable()->constrained('users')->nullOnDelete();
            $table->foreignId('locked_by')->nullable()->constrained('users')->nullOnDelete();
            $table->foreignUuid('category_id')->constrained('categories')->restrictOnDelete();
            $table->string('title');
            $table->string('slug')->unique();
            $table->longText('content');
            $table->enum('status', ['draft','review','scheduled','published','archived'])
                  ->default('draft')->index();
            $table->boolean('is_featured')->default(false);
            $table->unsignedBigInteger('view_count')->default(0)->index();
            $table->string('featured_image_url')->nullable();
            $table->timestamp('published_at')->nullable()->index();
            $table->timestamps();
            $table->softDeletes();
        });
    }

    public function down(): void
    {
        Schema::dropIfExists('articles');
    }
};''',
        
        "Potongan Kode 13. UserSeeder": '''// database/seeders/UserSeeder.php
class UserSeeder extends Seeder
{
    public function run(): void
    {
        $data = $this->loadJson();

        foreach ($data['users'] as $user) {
            DB::table('users')->insert([
                'name'       => $user['name'],
                'email'      => $user['email'],
                'password'   => Hash::make('password123'),
                'role'       => $user['role'],
                'is_active'  => $user['is_active'],
                'avatar_url' => $user['avatar_url'],
                'created_at' => now(),
                'updated_at' => now(),
            ]);
        }
    }

    private function loadJson(): array
    {
        return json_decode(file_get_contents(database_path('data/dummy2.json')), true);
    }
}''',
        
        "Potongan Kode 14. Artisan Command": '''// app/Console/Commands/PublishScheduledArticles.php
class PublishScheduledArticles extends Command
{
    protected $signature = 'articles:publish-scheduled';
    protected $description = 'Publish semua artikel terjadwal yang sudah waktunya tayang';

    public function __construct(private readonly ScheduledPublishService $service)
    {
        parent::__construct();
    }

    public function handle(): int
    {
        $this->info('[ScheduledPublish] Memulai proses auto-publish...');
        $count = $this->service->run();

        if ($count === 0) {
            $this->line('  → Tidak ada artikel yang perlu dipublish.');
        } else {
            $this->info("  → {$count} artikel berhasil dipublish dan diindeks.");
        }
        return Command::SUCCESS;
    }
}''',
        
        "Potongan Kode 15. Definisi TypeScript interfaces": '''// types/index.ts
export type Role = 'reader' | 'journalist' | 'editor' | 'admin';
export type ArticleStatus = 'draft' | 'review' | 'scheduled' | 'published' | 'archived';

export interface User {
  id: string;
  name: string;
  email: string;
  avatar_url?: string;
  role: Role;
  is_active: boolean;
}

export interface Article {
  id: string;
  title: string;
  slug: string;
  content: string;
  featured_image_url: string | null;
  status: ArticleStatus;
  category: Category;
  tags: Tag[];
  author: Pick<User, 'id' | 'name' | 'avatar_url'>;
  published_at: string | null;
}

export interface Bookmark {
  id: string;
  article: Pick<Article, 'id' | 'title' | 'slug' | 'featured_image_url'>;
  created_at: string;
}

export interface ApiSuccess<T> {
  status: 'success';
  data: T;
}

export interface ApiError {
  status: 'error';
  code: number;
  error: string;
  message: string;
}''',
        
        "Potongan Kode 16. Schema validasi Zod": '''// lib/validations.ts
import { z } from 'zod';

export const loginSchema = z.object({
  email: z.string().email('Format email tidak valid'),
  password: z.string().min(8, 'Password minimal 8 karakter'),
});

export const registerSchema = z.object({
  name: z.string().min(2, 'Nama minimal 2 karakter'),
  email: z.string().email('Format email tidak valid'),
  password: z.string().min(8, 'Password minimal 8 karakter'),
});

export const articleSchema = z.object({
  title: z.string().min(5, 'Judul minimal 5 karakter').max(255),
  content: z.string().min(1, 'Konten tidak boleh kosong'),
  category_id: z.string().uuid('Pilih kategori'),
  tags: z.array(z.string().uuid()).optional(),
});

export const scheduleSchema = z.object({
  scheduled_at: z.string().refine((val) => {
    const diff = new Date(val).getTime() - Date.now();
    return diff >= 5 * 60 * 1000;
  }, 'Waktu tayang minimal 5 menit dari sekarang'),
});

export const commentSchema = z.object({
  content: z.string().min(3, 'Komentar minimal 3 karakter')
                .max(1000, 'Komentar maksimal 1000 karakter'),
  parent_id: z.string().uuid().optional(),
});''',
        
        "Potongan Kode 17. Konstanta konfigurasi frontend": '''// constants/roles.ts
export const ROLES = {
  READER: 'reader',
  JOURNALIST: 'journalist',
  EDITOR: 'editor',
  ADMIN: 'admin',
} as const;

export const canPublish = (role: Role) =>
  ([ROLES.EDITOR, ROLES.ADMIN] as Role[]).includes(role);

export const canManageUsers = (role: Role) => role === ROLES.ADMIN;

// constants/errorMessages.ts
export const ERROR_MESSAGES: Record<string, string> = {
  INVALID_CREDENTIALS: 'Email atau password salah',
  TOKEN_EXPIRED: 'Sesi telah berakhir, silakan login ulang',
  FORBIDDEN: 'Anda tidak memiliki akses untuk aksi ini',
  ARTICLE_NOT_FOUND: 'Artikel tidak ditemukan',
  COMMENT_TOO_SHORT: 'Komentar minimal 3 karakter',
  COMMENT_RATE_LIMIT: 'Terlalu banyak komentar, coba lagi nanti',
  CANNOT_MODIFY_SELF: 'Tidak bisa mengubah akun sendiri',
};''',
        
        "Potongan Kode 18. Fungsi API frontend untuk CRUD artikel": '''// lib/api/articles.ts
import axiosInstance from '@/lib/axios';
import { ApiSuccess, Article, ArticleStatus, Pagination } from '@/app/types';

export const getArticles = (params?: {
  status?: ArticleStatus; category?: string; tag?: string;
  page?: number; limit?: number;
}) => axiosInstance.get<ApiSuccess<{ articles: Article[]; pagination: Pagination }>>(
  '/articles', { params }
);

export const getCmsArticles = () =>
  axiosInstance.get<ApiSuccess<Article[]>>('/cms/articles');

export const getArticleBySlug = (slug: string) =>
  axiosInstance.get<ApiSuccess<Article>>(`/articles/${slug}`);

export const createArticle = (data: {
  title: string; content: string; category_id: string;
  tags?: string[]; featured_image_url?: string;
}) => axiosInstance.post<ApiSuccess<{ id: string; slug: string }>>(
  '/cms/articles', data
);

export const updateArticleStatus = (id: string, data: {
  status: ArticleStatus; scheduled_at?: string; change_note?: string;
}) => axiosInstance.patch<ApiSuccess<{ id: string; status: ArticleStatus }>>(
  `/cms/articles/${id}/status`, data
);

export const lockArticle = (id: string) =>
  axiosInstance.post<ApiSuccess<null>>(`/cms/articles/${id}/lock`);

export const unlockArticle = (id: string) =>
  axiosInstance.post<ApiSuccess<null>>(`/cms/articles/${id}/unlock`);''',
        
        "Potongan Kode 19. Implementasi Zustand Auth Store": '''// stores/authStore.ts
import { create } from 'zustand';
import { persist } from 'zustand/middleware';
import { Role } from '@/app/types';

interface AuthUser {
  id: string;
  name: string;
  email?: string;
  role: Role;
}

interface AuthState {
  accessToken: string | null;
  user: AuthUser | null;
  isAuthenticated: boolean;
  setAuth: (token: string, user: AuthUser) => void;
  setAccessToken: (token: string) => void;
  logout: () => void;
}

export const useAuthStore = create<AuthState>()(
  persist(
    (set) => ({
      accessToken: null,
      user: null,
      isAuthenticated: false,

      setAuth: (token, user) =>
        set({ accessToken: token, user, isAuthenticated: true }),

      setAccessToken: (token) => set({ accessToken: token }),

      logout: () =>
        set({ accessToken: null, user: null, isAuthenticated: false }),
    }),
    { name: 'auth-storage' }
  )
);

export const getAccessToken = () => useAuthStore.getState().accessToken;
export const setAccessToken = (token: string) =>
  useAuthStore.getState().setAccessToken(token);
export const logoutStore = () => useAuthStore.getState().logout();''',
    }
    
    # Process paragraphs
    paragraphs = doc.paragraphs
    
    # Find each code snippet label and fill the paragraph immediately before it
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        
        # Check if this paragraph is a code snippet label
        for label_key, code in code_snippets.items():
            if label_key in text:
                # Look for the first empty paragraph before this label
                # within a reasonable range (max 10 paragraphs back)
                for j in range(max(0, i-10), i):
                    prev_para = paragraphs[j]
                    prev_text = prev_para.text.strip()
                    
                    # If we find an empty paragraph, fill it with the first line of code
                    if prev_text == '':
                        # Clear any existing runs
                        for run in prev_para.runs:
                            run.text = ''
                        # Add the code
                        set_code_format(prev_para, code)
                        break
                break
    
    return doc

def main():
    print("Opening FIX.docx...")
    doc = docx.Document('FIX.docx')
    
    print("Filling empty code snippets...")
    doc = fill_code_snippets_v2(doc)
    
    print("Saving modified document as FIX_filled.docx...")
    doc.save('FIX_filled.docx')
    
    print("Done! Code snippets have been filled.")
    print("\nNote: Diagram placeholders (Gambar 3.1 - 3.18) require actual image files.")
    print("The script has filled all empty code snippets (Potongan Kode 1-19).")

if __name__ == '__main__':
    main()
