"""
Script untuk memodifikasi laporan PLK Ismy:
- Menambahkan "Preview Berita" ke daftar halaman CMS
- Menambahkan mention preview berita di bagian implementasi, testing, dan rekognisi
"""
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

INPUT  = r'c:\apps\klojen.com\_PLK_ismy .docx'
OUTPUT = r'c:\apps\klojen.com\_PLK_ismy_updated.docx'

doc = Document(INPUT)

def get_text(p):
    """Get full text of a paragraph."""
    return ''.join(r.text for r in p.runs)

def find_para(search, start=0):
    """Find paragraph index containing search text."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if search in get_text(p):
            return i
    return -1

def insert_paragraph_after(idx, text, style=None):
    """Insert a new paragraph after the paragraph at idx."""
    ref = doc.paragraphs[idx]
    new_p = copy.deepcopy(ref._element)
    # Clear existing runs
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    # Add new run with text
    new_run = copy.deepcopy(ref._element.findall(qn('w:r'))[0]) if ref._element.findall(qn('w:r')) else None
    if new_run is None:
        from lxml import etree
        new_run = etree.SubElement(new_p, qn('w:r'))
    # Set text
    t_elem = new_run.find(qn('w:t'))
    if t_elem is None:
        from lxml import etree
        t_elem = etree.SubElement(new_run, qn('w:t'))
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    # Insert after reference
    ref._element.addnext(new_p)
    return new_p

changes = 0

# ============================================================
# 1. III.2.1 - Tambah "Preview Berita" ke daftar CMS pages
# ============================================================
idx = find_para('Tinjauan Artikel')
if idx >= 0:
    t = get_text(doc.paragraphs[idx])
    if 'Preview Berita' not in t and 'preview berita' not in t:
        # Modify the Tinjauan Artikel paragraph to add Preview Berita after it
        # Find "Profil & Edit Profil" paragraph and insert before it
        idx_profil = find_para('Profil & Edit Profil', idx)
        if idx_profil >= 0:
            # Insert new paragraph for Preview Berita before Profil
            p_elem = doc.paragraphs[idx_profil]._element
            from lxml import etree
            new_p = copy.deepcopy(doc.paragraphs[idx]._element)
            # Clear runs
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            # Create new run
            new_r = copy.deepcopy(doc.paragraphs[idx]._element.findall(qn('w:r'))[0])
            t_elem = new_r.find(qn('w:t'))
            if t_elem is not None:
                t_elem.text = 'Preview Berita: Halaman untuk melihat tampilan artikel sebelum dipublikasikan, menampilkan konten lengkap beserta gambar, tag, dan metadata artikel dalam format yang menyerupai tampilan portal berita publik.'
                t_elem.set(qn('xml:space'), 'preserve')
            new_p.append(new_r)
            p_elem.addprevious(new_p)
            changes += 1
            print(f"[1] Added Preview Berita paragraph before Profil (para {idx_profil})")

# ============================================================
# 2. III.2.1 - Tambah "Preview Berita" ke portal berita list
#    (Actually, we want to add it to CMS list, not portal list)
#    The current structure already has portal pages separate.
#    Let's fix the numbering: add "5." to Preview Berita
# ============================================================

# ============================================================
# 3. III.2.2 - Add mention of preview berita page
# ============================================================
idx = find_para('Halaman lain yang dikembangkan meliputi halaman daftar artikel')
if idx >= 0:
    t = get_text(doc.paragraphs[idx])
    if 'preview berita' not in t.lower():
        # Add "halaman preview berita" to the list
        new_text = t.replace(
            'serta halaman edit profil',
            'halaman preview berita untuk melihat tampilan artikel sebelum publish, serta halaman edit profil'
        )
        if new_text != t:
            for run in doc.paragraphs[idx].runs:
                if 'serta halaman edit profil' in run.text:
                    run.text = run.text.replace(
                        'serta halaman edit profil',
                        'halaman preview berita untuk melihat tampilan artikel sebelum publish, serta halaman edit profil'
                    )
                    changes += 1
                    print(f"[3] Updated III.2.2 paragraph to include preview berita")
                    break

# ============================================================
# 4. III.3 Pencapaian - Add preview berita to CMS page list
# ============================================================
idx = find_para('Desain UI/UX Berhasil Dirancang')
if idx >= 0:
    t = get_text(doc.paragraphs[idx])
    if 'preview berita' not in t.lower():
        new_text = t.replace(
            'profil, beranda portal berita',
            'profil, preview berita, beranda portal berita'
        )
        if new_text != t:
            for run in doc.paragraphs[idx].runs:
                if 'profil, beranda portal berita' in run.text:
                    run.text = run.text.replace(
                        'profil, beranda portal berita',
                        'profil, preview berita, beranda portal berita'
                    )
                    changes += 1
                    print(f"[4] Updated Pencapaian - Desain UI/UX")
                    break

idx = find_para('Frontend CMS Berfungsi Penuh')
if idx >= 0:
    t = get_text(doc.paragraphs[idx])
    if 'preview berita' not in t.lower():
        new_text = t.replace(
            'tinjauan artikel, serta pengelolaan profil',
            'tinjauan artikel, preview berita, serta pengelolaan profil'
        )
        if new_text != t:
            for run in doc.paragraphs[idx].runs:
                if 'tinjauan artikel, serta pengelolaan profil' in run.text:
                    run.text = run.text.replace(
                        'tinjauan artikel, serta pengelolaan profil',
                        'tinjauan artikel, preview berita, serta pengelolaan profil'
                    )
                    changes += 1
                    print(f"[4b] Updated Pencapaian - Frontend CMS")
                    break

# ============================================================
# 5. III.4.3 Testing CMS - Add test row for Preview Berita
# ============================================================
# Find the table for Testing Halaman CMS (III.4.3)
idx_heading = find_para('III.4.3 Testing Halaman CMS')
if idx_heading >= 0:
    # Find the table after this heading
    # Tables are separate from paragraphs, so we look through document tables
    for table in doc.tables:
        # Check if this table has "Responsivitas CMS" (last row of CMS test table)
        for row in table.rows:
            cells_text = [c.text.strip() for c in row.cells]
            if any('Responsivitas CMS' in ct for ct in cells_text):
                # Check if preview berita already in table
                table_text = '\n'.join(c.text for row in table.rows for c in row.cells)
                if 'preview berita' not in table_text.lower() and 'Preview Berita' not in table_text:
                    # Find the row with "Responsivitas CMS" and change its number to 6
                    # Then insert a new row 5 for Preview Berita
                    for ri, row in enumerate(table.rows):
                        cells_text = [c.text.strip() for c in row.cells]
                        if any('Responsivitas CMS' in ct for ct in cells_text):
                            # Change the number from 5 to 6
                            for c in row.cells:
                                if c.text.strip() == '5.':
                                    for p in c.paragraphs:
                                        for r in p.runs:
                                            if '5' in r.text:
                                                r.text = r.text.replace('5', '6')
                                                break
                                    break
                            
                            # Insert new row before this one
                            new_row = copy.deepcopy(row._tr)
                            # Clear cells content and set new values
                            cell_data = [
                                '5.',
                                'Buka Preview Berita',
                                'Klik salah satu artikel, buka halaman Preview Berita.',
                                'Tampil preview artikel lengkap (konten, gambar, tag, metadata) menyerupai tampilan portal berita.',
                                ''
                            ]
                            for ci, cell in enumerate(new_row.findall(qn('w:tc'))):
                                # Clear existing paragraphs
                                for p in cell.findall(qn('w:p')):
                                    for r in p.findall(qn('w:r')):
                                        t = r.find(qn('w:t'))
                                        if t is not None:
                                            t.text = cell_data[ci] if ci < len(cell_data) else ''
                                            t.set(qn('xml:space'), 'preserve')
                                            # Keep only first run
                                            break
                                    # Remove extra runs
                                    runs = p.findall(qn('w:r'))
                                    for extra_r in runs[1:]:
                                        p.remove(extra_r)
                            
                            row._tr.addprevious(new_row)
                            changes += 1
                            print(f"[5] Added Preview Berita test row to CMS testing table")
                            break
                break

# ============================================================
# 6. IV.3 Rekognisi Uji Coba - Add preview berita
# ============================================================
idx = find_para('pengujian halaman CMS, meliputi filter artikel')
if idx >= 0:
    t = get_text(doc.paragraphs[idx])
    if 'preview berita' not in t.lower():
        new_text = t.replace(
            'tinjauan editor, dan responsivitas',
            'tinjauan editor, preview berita, dan responsivitas'
        )
        if new_text != t:
            for run in doc.paragraphs[idx].runs:
                if 'tinjauan editor, dan responsivitas' in run.text:
                    run.text = run.text.replace(
                        'tinjauan editor, dan responsivitas',
                        'tinjauan editor, preview berita, dan responsivitas'
                    )
                    changes += 1
                    print(f"[6] Updated Rekognisi Uji Coba - CMS testing")
                    break

# Also update the implementasi paragraph in rekognisi
idx = find_para('halaman-halaman CMS (daftar artikel, tulis berita, tinjauan artikel, dan profil)')
if idx >= 0:
    t = get_text(doc.paragraphs[idx])
    if 'preview berita' not in t.lower():
        new_text = t.replace(
            'tinjauan artikel, dan profil',
            'tinjauan artikel, preview berita, dan profil'
        )
        if new_text != t:
            for run in doc.paragraphs[idx].runs:
                if 'tinjauan artikel, dan profil' in run.text:
                    run.text = run.text.replace(
                        'tinjauan artikel, dan profil',
                        'tinjauan artikel, preview berita, dan profil'
                    )
                    changes += 1
                    print(f"[6b] Updated Rekognisi Implementasi - CMS pages")
                    break

# ============================================================
# 7. V.2 Saran - Remove "Penambahan fitur preview artikel" since it already exists
# ============================================================
idx = find_para('Penambahan fitur preview artikel sebelum dipublikasikan')
if idx >= 0:
    t = get_text(doc.paragraphs[idx])
    # Change to a different suggestion since preview already exists
    new_text = t.replace(
        'Penambahan fitur preview artikel sebelum dipublikasikan agar editor dapat melihat tampilan akhir artikel.',
        'Penyempurnaan fitur preview berita dengan penambahan opsi preview pada berbagai ukuran layar (desktop, tablet, mobile).'
    )
    if new_text != t:
        for run in doc.paragraphs[idx].runs:
            if 'Penambahan fitur preview artikel' in run.text:
                run.text = run.text.replace(
                    'Penambahan fitur preview artikel sebelum dipublikasikan agar editor dapat melihat tampilan akhir artikel.',
                    'Penyempurnaan fitur preview berita dengan penambahan opsi preview pada berbagai ukuran layar (desktop, tablet, mobile).'
                )
                changes += 1
                print(f"[7] Updated Saran - changed preview suggestion")
                break

# ============================================================
# Save
# ============================================================
doc.save(OUTPUT)
print(f"\n=== Done! {changes} changes applied ===")
print(f"Output: {OUTPUT}")
